#!/usr/bin/env python3
"""Generate UCF Business Course — Resource Guide Word Document"""
import sys, io, os, json, glob, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'], capture_output=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))

# ─── Color palette ───
UCF_GOLD = RGBColor(0xFC, 0xB5, 0x14)
UCF_BLACK = RGBColor(0x0F, 0x0F, 0x0F)
SECTION_BG = RGBColor(0xF5, 0xF5, 0xF5)
GREEN = RGBColor(0x27, 0xAE, 0x60)
BLUE = RGBColor(0x1A, 0x73, 0xE8)
GRAY = RGBColor(0x55, 0x65, 0x75)

def set_cell_bg(cell, hex_color):
    """Set table cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p

def add_para(doc, text, bold=False, italic=False, size=10, color=None, left_indent=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    return p

# ─── RESOURCE DATA ───
# Organized by lesson, then by category for the master list

MODULES = {
    0: {"location": "Highway", "title": "Start Here — U.S. Market Overview"},
    1: {"location": "Downtown", "title": "Legal Setup"},
    2: {"location": "Financial District", "title": "Taxes & Banking"},
    3: {"location": "Space Center", "title": "IP, Insurance & Contracts"},
    4: {"location": "Theme Park District", "title": "Hiring & Compliance"},
    5: {"location": "Harbor", "title": "Product & Market"},
    6: {"location": "Beach", "title": "U.S. Business Culture"},
    7: {"location": "Natural Springs", "title": "Central Florida Ecosystem"},
    8: {"location": "UCF Campus", "title": "Readiness & Next Steps"},
}

RESOURCES_BY_LESSON = {
    "lesson-1-1": [
        {"name": "SunBiz — Florida Division of Corporations", "url": "https://sunbiz.org", "type": "Gov Tool", "cost": "$125 LLC / $70 Corp", "for": "LLC or C-Corp registration and name search", "language": "EN"},
        {"name": "IRS EIN Application", "url": "https://www.irs.gov/businesses/small-businesses-self-employed/employer-id-numbers", "type": "Gov Service", "cost": "Free", "for": "Federal tax ID number (required for banking)", "language": "EN"},
    ],
    "lesson-1-2": [
        {"name": "SunBiz — Florida Division of Corporations", "url": "https://sunbiz.org", "type": "Gov Tool", "cost": "$125 LLC / $70 Corp", "for": "Step-by-step business registration in Florida", "language": "EN"},
    ],
    "lesson-2-1": [
        {"name": "IRS — Small Business Tax Center", "url": "https://www.irs.gov/businesses/small-businesses-self-employed", "type": "Gov Agency", "cost": "Free", "for": "Federal tax guides, quarterly payment info, entity types", "language": "EN"},
    ],
    "lesson-2-3": [
        {"name": "Mercury", "url": "https://mercury.com", "type": "FinTech", "cost": "Free (no min balance)", "for": "Online business banking, friendly for international founders", "language": "EN"},
        {"name": "Relay", "url": "https://relayfi.com", "type": "FinTech", "cost": "Free", "for": "Business banking with accounting software integration", "language": "EN"},
        {"name": "Stripe Atlas", "url": "https://stripe.com/atlas", "type": "Commercial", "cost": "$500 setup", "for": "Company formation + banking bundle for international founders", "language": "EN"},
    ],
    "lesson-3-1": [
        {"name": "USPTO — U.S. Patent and Trademark Office", "url": "https://www.uspto.gov", "type": "Gov Agency", "cost": "$250–$350/class", "for": "Trademark registration, patent filings, IP search", "language": "EN"},
    ],
    "lesson-3-3": [
        {"name": "Florida Bar Lawyer Referral Service", "url": "https://www.floridabar.org/public/lrs/", "type": "Bar Association", "cost": "$25 referral fee", "for": "Find a vetted Florida business attorney", "language": "EN"},
        {"name": "LegalZoom", "url": "https://www.legalzoom.com", "type": "Commercial", "cost": "Varies", "for": "Formation and basic legal document services", "language": "EN"},
    ],
    "lesson-4-2": [
        {"name": "USCIS — U.S. Citizenship and Immigration Services", "url": "https://www.uscis.gov", "type": "Gov Agency", "cost": "Free (info)", "for": "Official visa classifications, work authorization, status", "language": "EN"},
        {"name": "AILA Lawyer Search", "url": "https://www.ailalawyer.com", "type": "Professional Referral", "cost": "Free search", "for": "Find AILA-member immigration attorneys", "language": "EN"},
        {"name": "Fragomen", "url": "https://www.fragomen.com", "type": "Law Firm", "cost": "Varies", "for": "International immigration law firm with FL presence", "language": "EN"},
    ],
    "lesson-7-1": [
        {"name": "Florida SBDC at UCF", "url": "https://sbdcorlando.com", "type": "Gov-Funded Nonprofit", "cost": "Free", "for": "Business consulting, market research, financial projections", "language": "EN/ES"},
        {"name": "SCORE Orlando", "url": "https://orlando.score.org", "type": "Nonprofit", "cost": "Free", "for": "Mentoring from retired executives, ongoing coaching", "language": "EN"},
        {"name": "SBA — Small Business Administration", "url": "https://www.sba.gov", "type": "Gov Agency", "cost": "Free", "for": "Loan programs, business training, certifications", "language": "EN/ES"},
        {"name": "Orlando Economic Partnership", "url": "https://orlando.org", "type": "Economic Dev Org", "cost": "Free", "for": "International business relocation and expansion support", "language": "EN"},
        {"name": "Prospera", "url": "https://prosperausa.org", "type": "Nonprofit", "cost": "Free/Low-cost", "for": "Bilingual business development for Hispanic entrepreneurs", "language": "EN/ES"},
        {"name": "UCF Business Incubation Program", "url": "https://incubator.ucf.edu", "type": "University Program", "cost": "Free (Phase 1)", "for": "Mentorship, workspace, advisory (Phases 2–3 after this course)", "language": "EN"},
        {"name": "Enterprise Florida", "url": "https://www.enterpriseflorida.com", "type": "State Org", "cost": "Free", "for": "State trade assistance and export development programs", "language": "EN"},
    ],
    "lesson-7-2": [
        {"name": "Prospera", "url": "https://prosperausa.org", "type": "Nonprofit", "cost": "Free/Low-cost", "for": "Workshops and 1-on-1 consulting in Spanish", "language": "EN/ES"},
        {"name": "Hispanic Chamber of Commerce of Metro Orlando", "url": "https://hispanicchamber.com", "type": "Chamber", "cost": "$200–$500/year", "for": "Business networking, bilingual workshops, advocacy", "language": "EN/ES"},
    ],
    "lesson-7-4": [
        {"name": "CFITO — Central Florida International Trade Office", "url": "https://cfito.com", "type": "Trade Org", "cost": "Free/Nominal", "for": "Trade delegations, market research, buyer/distributor intros", "language": "EN"},
        {"name": "Prospera", "url": "https://prosperausa.org", "type": "Nonprofit", "cost": "Free/Low-cost", "for": "Hispanic/LATAM founder support — first point of contact", "language": "EN/ES"},
        {"name": "Hispanic Chamber of Commerce of Metro Orlando", "url": "https://hispanicchamber.com", "type": "Chamber", "cost": "$200–$500/year", "for": "LATAM business networking and partner connections", "language": "EN/ES"},
    ],
    "lesson-7-5": [
        {"name": "Prospera", "url": "https://prosperausa.org", "type": "Nonprofit", "cost": "Free/Low-cost", "for": "Mentoring, business training, micro-lending for Hispanic founders", "language": "EN/ES"},
        {"name": "SCORE Orlando", "url": "https://orlando.score.org", "type": "Nonprofit", "cost": "Free", "for": "Industry-matched mentor for ongoing coaching", "language": "EN"},
        {"name": "CFITO — Central Florida International Trade Office", "url": "https://cfito.com", "type": "Trade Org", "cost": "Free/Nominal", "for": "Import/export, LATAM-Caribbean trade corridor support", "language": "EN"},
        {"name": "Hispanic Chamber of Commerce of Metro Orlando", "url": "https://hispanicchamber.com", "type": "Chamber", "cost": "$200–$500/year", "for": "Events, bilingual connections, first local partner referrals", "language": "EN/ES"},
        {"name": "UCF BIP — Schedule 1:1 Meeting", "url": "https://calendly.com/br770425-ucf/30min", "type": "Scheduling", "cost": "Free", "for": "Book your 1:1 consultation with the UCF BIP team", "language": "EN"},
    ],
    "lesson-8-2": [
        {"name": "SunBiz — Florida Division of Corporations", "url": "https://sunbiz.org", "type": "Gov Tool", "cost": "$125 LLC / $70 Corp", "for": "Entity name check and registration", "language": "EN"},
        {"name": "IRS EIN Application", "url": "https://www.irs.gov/businesses/small-businesses-self-employed/employer-id-numbers", "type": "Gov Service", "cost": "Free", "for": "Employer Identification Number", "language": "EN"},
        {"name": "Florida Bar Lawyer Referral Service", "url": "https://www.floridabar.org/public/lrs/", "type": "Bar Association", "cost": "$25 referral fee", "for": "Vetted Florida business attorney search", "language": "EN"},
        {"name": "LegalZoom", "url": "https://www.legalzoom.com", "type": "Commercial", "cost": "Varies", "for": "Formation and legal services", "language": "EN"},
        {"name": "Incfile", "url": "https://www.incfile.com", "type": "Commercial", "cost": "Varies", "for": "Business entity formation service", "language": "EN"},
        {"name": "Mercury", "url": "https://mercury.com", "type": "FinTech", "cost": "Free", "for": "Startup-friendly online banking", "language": "EN"},
        {"name": "Relay", "url": "https://relayfi.com", "type": "FinTech", "cost": "Free", "for": "Business banking + accounting integration", "language": "EN"},
        {"name": "Stripe Atlas", "url": "https://stripe.com/atlas", "type": "Commercial", "cost": "$500 setup", "for": "Formation + banking bundle", "language": "EN"},
        {"name": "QuickBooks Online", "url": "https://quickbooks.intuit.com", "type": "Software", "cost": "$15–$50/mo", "for": "Small business accounting — most popular in U.S.", "language": "EN"},
        {"name": "USCIS", "url": "https://www.uscis.gov", "type": "Gov Agency", "cost": "Free (info)", "for": "Official immigration status and visa info", "language": "EN"},
        {"name": "Boundless", "url": "https://www.boundless.com", "type": "Immigration Service", "cost": "Varies", "for": "Immigration application assistance", "language": "EN"},
        {"name": "AILA Lawyer Search", "url": "https://www.ailalawyer.com", "type": "Professional Referral", "cost": "Free", "for": "Immigration attorney finder", "language": "EN"},
        {"name": "Fragomen", "url": "https://www.fragomen.com", "type": "Law Firm", "cost": "Varies", "for": "International immigration legal services", "language": "EN"},
        {"name": "Florida SBDC at UCF", "url": "https://sbdcorlando.com", "type": "Gov-Funded Nonprofit", "cost": "Free", "for": "Free consulting and market research", "language": "EN/ES"},
        {"name": "SCORE Orlando", "url": "https://orlando.score.org", "type": "Nonprofit", "cost": "Free", "for": "Business mentoring", "language": "EN"},
        {"name": "Orlando Economic Partnership", "url": "https://orlando.org", "type": "Economic Dev", "cost": "Free", "for": "Regional expansion support", "language": "EN"},
        {"name": "Prospera", "url": "https://prosperausa.org", "type": "Nonprofit", "cost": "Free/Low-cost", "for": "Hispanic founder support", "language": "EN/ES"},
        {"name": "UCF Business Incubation Program", "url": "https://incubator.ucf.edu", "type": "University Program", "cost": "Free (Phase 1)", "for": "Mentorship, advisory, phases 2–3", "language": "EN"},
    ],
    "lesson-8-3": [
        {"name": "UCF BIP — Book Your 1:1 Meeting", "url": "https://calendly.com/br770425-ucf/30min", "type": "Scheduling", "cost": "Free", "for": "30-min consultation with UCF BIP team (Brian Bedrick)", "language": "EN"},
        {"name": "Florida SBDC at UCF", "url": "https://sbdcorlando.com", "type": "Gov-Funded Nonprofit", "cost": "Free", "for": "Pre-meeting counseling session recommended", "language": "EN/ES"},
        {"name": "UCF Business Incubation Program", "url": "https://incubator.ucf.edu", "type": "University Program", "cost": "Varies by phase", "for": "Phase 2 and Phase 3 market entry program details", "language": "EN"},
    ],
}

LESSON_TITLES = {
    "lesson-0-1": "U.S. Market Entry — What You're About to Learn",
    "lesson-1-1": "Choosing Your U.S. Business Entity",
    "lesson-1-2": "Registering Your Business in Florida",
    "lesson-1-3": "Opening Your Business Bank Account",
    "lesson-2-1": "Federal and State Taxes — How the System Works",
    "lesson-2-2": "Sales Tax, Accounting, and Bookkeeping",
    "lesson-2-3": "U.S. Banking — When Banks Say No",
    "lesson-3-1": "Trademarks, Patents, and Copyright",
    "lesson-3-2": "Business Insurance — What You Actually Need",
    "lesson-3-3": "Contracts and Legal Help",
    "lesson-4-1": "Your First U.S. Hire — Contractor vs. Employee",
    "lesson-4-2": "Visa and Work Authorization for Founders",
    "lesson-4-3": "Payroll, Benefits, and HR Compliance",
    "lesson-5-1": "Adapting Your Product for the U.S. Market",
    "lesson-5-2": "Pricing, Sales, and Vendor Onboarding",
    "lesson-6-1": "How Americans Communicate in Business",
    "lesson-6-2": "Networking and Business Relationships",
    "lesson-6-3": "Cultural Differences and Negotiation",
    "lesson-7-1": "Your Central Florida Support Network",
    "lesson-7-2": "Central Florida's Business Ecosystem",
    "lesson-7-3": "Central Florida by the Numbers",
    "lesson-7-4": "Latin American Founders and the LATAM Trade Corridor",
    "lesson-7-5": "UCF BIP Programs — Your Next Step",
    "lesson-8-1": "Self-Assessment — Where Do You Stand?",
    "lesson-8-2": "Resource Downloads and Digital Credential",
    "lesson-8-3": "Your Next Steps with UCF BIP",
}

def get_module_for_lesson(lid):
    parts = lid.split('-')
    if len(parts) >= 2:
        m = parts[1]
        try:
            return int(m)
        except:
            return 0
    return 0

# ─── BUILD DOCUMENT ───
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ─── COVER PAGE ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UCF Business Incubation Program')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = GRAY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Phase 1 Program — Complete Resource Guide')
run2.bold = True
run2.font.size = Pt(22)
run2.font.color.rgb = UCF_BLACK

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    'All external links, tools, organizations, and services referenced in the course,\n'
    'organized by lesson and by category. For review and validation — April 2026.'
)
run3.font.size = Pt(11)
run3.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_page_break()

# ─── SECTION 1: HOW TO USE THIS DOCUMENT ───
add_heading(doc, 'How to Use This Document', level=1)
add_para(doc,
    'This guide catalogs every external resource, link, tool, and organization referenced in the UCF Business Phase 1 Program. '
    'It serves three purposes:',
    size=10)
bullets = [
    'Content review — confirm every linked resource is current, accurate, and appropriate for the audience',
    'Legal/compliance review — verify that linked commercial services and government tools are correctly described',
    'Translation check — identify which resources offer Spanish language support (marked EN/ES)',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()
add_para(doc,
    'Resources are organized two ways: (1) by module and lesson — showing exactly where each resource appears in the course, '
    'and (2) by category — a master reference table for quick lookup.',
    size=10)
doc.add_paragraph()

# ─── SECTION 2: RESOURCES BY MODULE AND LESSON ───
add_heading(doc, 'Resources by Module and Lesson', level=1)
add_para(doc,
    'The following sections list all resources in the order a learner encounters them as they progress through the course.',
    size=10, color=GRAY)
doc.add_paragraph()

current_module = -1
for lid in sorted(RESOURCES_BY_LESSON.keys()):
    resources = RESOURCES_BY_LESSON[lid]
    if not resources:
        continue
    mod_num = get_module_for_lesson(lid)
    mod_info = MODULES.get(mod_num, {})
    lesson_title = LESSON_TITLES.get(lid, lid)

    if mod_num != current_module:
        current_module = mod_num
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"MODULE {mod_num} — {mod_info.get('location','').upper()}: {mod_info.get('title','').upper()}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = UCF_GOLD
        p.paragraph_format.space_after = Pt(4)

    # Lesson subheader
    p = doc.add_paragraph()
    run = p.add_run(f'{lid.upper()} — {lesson_title}')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = UCF_BLACK
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    # Resource table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0].cells
    headers = ['Resource', 'URL', 'Cost', 'EN/ES']
    widths = [Inches(2.2), Inches(2.5), Inches(0.8), Inches(0.5)]
    for i, (h, w) in enumerate(zip(headers, widths)):
        hdr[i].width = w
        set_cell_bg(hdr[i], 'F0F0F0')
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows
    for res in resources:
        row = table.add_row().cells
        row[0].width = widths[0]
        row[1].width = widths[1]
        row[2].width = widths[2]
        row[3].width = widths[3]

        # Resource name + for
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(res['name'])
        r0.bold = True
        r0.font.size = Pt(9)
        p0b = row[0].add_paragraph(res.get('for',''))
        p0b.runs[0].font.size = Pt(8)
        p0b.runs[0].font.color.rgb = GRAY

        # URL
        p1 = row[1].paragraphs[0]
        r1 = p1.add_run(res.get('url','—'))
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = BLUE

        # Cost
        p2 = row[2].paragraphs[0]
        r2 = p2.add_run(res.get('cost','—'))
        r2.font.size = Pt(9)

        # Language
        lang = res.get('language','EN')
        p3 = row[3].paragraphs[0]
        r3 = p3.add_run(lang)
        r3.font.size = Pt(9)
        if 'ES' in lang:
            r3.font.color.rgb = GREEN
            r3.bold = True

    doc.add_paragraph()

doc.add_page_break()

# ─── SECTION 3: MASTER RESOURCE TABLE ───
add_heading(doc, 'Master Resource Reference', level=1)
add_para(doc,
    'All resources alphabetically by category. Use this table to verify links, costs, and descriptions before distribution.',
    size=10, color=GRAY)
doc.add_paragraph()

CATEGORIES = [
    ("Government Agencies & Tools", [
        {"name": "IRS — Small Business Tax Center", "url": "https://www.irs.gov/businesses/small-businesses-self-employed", "cost": "Free", "for": "Federal tax guides, EIN, quarterly payments", "lessons": "1-1, 2-1, 8-2", "lang": "EN"},
        {"name": "IRS EIN Application", "url": "https://www.irs.gov/businesses/small-businesses-self-employed/employer-id-numbers", "cost": "Free", "for": "Get your federal Employer Identification Number", "lessons": "1-1, 8-2", "lang": "EN"},
        {"name": "SBA — Small Business Administration", "url": "https://www.sba.gov", "cost": "Free", "for": "Loan programs, training, certifications", "lessons": "7-1", "lang": "EN/ES"},
        {"name": "SunBiz — Florida Division of Corporations", "url": "https://sunbiz.org", "cost": "$125 LLC / $70 Corp", "for": "Register your Florida business entity", "lessons": "1-1, 1-2, 3-3, 8-2", "lang": "EN"},
        {"name": "USPTO — Patent & Trademark Office", "url": "https://www.uspto.gov", "cost": "$250–$350/class", "for": "Trademark registration, patent search & filing", "lessons": "3-1, 3-3, 8-1, 8-2", "lang": "EN"},
        {"name": "USCIS — Immigration Services", "url": "https://www.uscis.gov", "cost": "Free (info)", "for": "Visa types, work authorization, immigration status", "lessons": "4-2, 8-2", "lang": "EN"},
    ]),
    ("Banking & Finance", [
        {"name": "Mercury", "url": "https://mercury.com", "cost": "Free", "for": "Online banking for startups, no min balance, international-friendly", "lessons": "2-3, 8-2", "lang": "EN"},
        {"name": "QuickBooks Online", "url": "https://quickbooks.intuit.com", "cost": "$15–$50/mo", "for": "Most popular U.S. small business accounting software", "lessons": "8-2", "lang": "EN"},
        {"name": "Relay", "url": "https://relayfi.com", "cost": "Free", "for": "Business banking with accounting integration", "lessons": "2-3, 8-2", "lang": "EN"},
        {"name": "Stripe Atlas", "url": "https://stripe.com/atlas", "cost": "$500 setup", "for": "Formation + banking bundle for international founders", "lessons": "2-3, 8-2", "lang": "EN"},
    ]),
    ("Legal & Formation Services", [
        {"name": "Florida Bar Lawyer Referral", "url": "https://www.floridabar.org/public/lrs/", "cost": "$25 referral fee", "for": "Find a vetted Florida business attorney", "lessons": "3-3, 8-2", "lang": "EN"},
        {"name": "Incfile", "url": "https://www.incfile.com", "cost": "Varies", "for": "Paid entity formation service", "lessons": "8-2", "lang": "EN"},
        {"name": "LegalZoom", "url": "https://www.legalzoom.com", "cost": "Varies", "for": "Paid formation and legal document services", "lessons": "3-3, 8-2", "lang": "EN"},
    ]),
    ("Immigration Services", [
        {"name": "AILA Lawyer Search", "url": "https://www.ailalawyer.com", "cost": "Free search", "for": "Find AILA-member immigration attorneys", "lessons": "4-2, 8-2", "lang": "EN"},
        {"name": "Boundless", "url": "https://www.boundless.com", "cost": "Varies", "for": "Immigration application assistance", "lessons": "8-2", "lang": "EN"},
        {"name": "Fragomen", "url": "https://www.fragomen.com", "cost": "Varies", "for": "International immigration law firm with FL office", "lessons": "4-2, 8-2", "lang": "EN"},
    ]),
    ("Nonprofit & Community Organizations (Central Florida)", [
        {"name": "CFITO — Central Florida International Trade Office", "url": "https://cfito.com", "cost": "Free / Nominal", "for": "Trade delegations, market research, LATAM partner intros", "lessons": "7-4, 7-5", "lang": "EN"},
        {"name": "Enterprise Florida", "url": "https://www.enterpriseflorida.com", "cost": "Free", "for": "State-level trade assistance and export development", "lessons": "7-1", "lang": "EN"},
        {"name": "Florida SBDC at UCF", "url": "https://sbdcorlando.com", "cost": "Free", "for": "Business consulting, market research, financial projections", "lessons": "7-1, 8-2, 8-3", "lang": "EN/ES"},
        {"name": "Hispanic Chamber of Commerce of Metro Orlando", "url": "https://hispanicchamber.com", "cost": "$200–$500/year", "for": "Bilingual networking, workshops, advocacy, partner referrals", "lessons": "7-2, 7-4, 7-5", "lang": "EN/ES"},
        {"name": "Orlando Economic Partnership", "url": "https://orlando.org", "cost": "Free", "for": "International business relocation and expansion support", "lessons": "7-1, 8-2", "lang": "EN"},
        {"name": "Prospera", "url": "https://prosperausa.org", "cost": "Free / Low-cost", "for": "Bilingual mentoring, training, micro-lending for Hispanic founders", "lessons": "4-2, 7-1, 7-2, 7-4, 7-5, 8-2", "lang": "EN/ES"},
        {"name": "SCORE Orlando", "url": "https://orlando.score.org", "cost": "Free", "for": "Free mentoring from retired executives, industry-matched", "lessons": "7-1, 7-5, 8-2", "lang": "EN"},
    ]),
    ("UCF Business Incubation Program", [
        {"name": "UCF BIP — Main Site", "url": "https://incubator.ucf.edu", "cost": "Free (Phase 1)", "for": "Program information, phases 2–3 market entry/soft landing", "lessons": "4-2, 7-1, 7-5, 8-1, 8-2, 8-3", "lang": "EN"},
        {"name": "UCF BIP — Book 1:1 Meeting", "url": "https://calendly.com/br770425-ucf/30min", "cost": "Free", "for": "Schedule your post-program consultation with Brian Bedrick", "lessons": "7-5, 8-3", "lang": "EN"},
    ]),
]

for cat_name, resources in CATEGORIES:
    p = doc.add_paragraph()
    run = p.add_run(cat_name)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = UCF_BLACK
    p.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    headers = ['Resource', 'URL', 'Cost', 'Appears In', 'EN/ES']
    widths = [Inches(1.8), Inches(2.0), Inches(0.7), Inches(1.0), Inches(0.5)]
    for i, (h, w) in enumerate(zip(headers, widths)):
        hdr[i].width = w
        set_cell_bg(hdr[i], 'E8E8E8')
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    for res in resources:
        row = table.add_row().cells
        for i, w in enumerate(widths):
            row[i].width = w

        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(res['name'])
        r0.bold = True
        r0.font.size = Pt(9)
        p0b = row[0].add_paragraph(res.get('for',''))
        p0b.runs[0].font.size = Pt(8)
        p0b.runs[0].font.color.rgb = GRAY

        p1 = row[1].paragraphs[0]
        r1 = p1.add_run(res.get('url',''))
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = BLUE

        p2 = row[2].paragraphs[0]
        r2 = p2.add_run(res.get('cost',''))
        r2.font.size = Pt(9)

        p3 = row[3].paragraphs[0]
        r3 = p3.add_run(f"Lessons: {res.get('lessons','')}")
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = GRAY

        lang = res.get('lang','EN')
        p4 = row[4].paragraphs[0]
        r4 = p4.add_run(lang)
        r4.font.size = Pt(9)
        if 'ES' in lang:
            r4.font.color.rgb = GREEN
            r4.bold = True

    doc.add_paragraph()

doc.add_page_break()

# ─── SECTION 4: LINKS FOR REVIEW ───
add_heading(doc, 'Links for Review', level=1)
add_para(doc,
    'The following items should be verified before the course is published. '
    'Please check that URLs are current, descriptions are accurate, and cost information is up to date.',
    size=10)
doc.add_paragraph()

review_items = [
    ("UCF BIP Calendly Link", "https://calendly.com/br770425-ucf/30min",
     "Verify this is the correct booking link for Brian Bedrick and it is active."),
    ("CFITO Website", "https://cfito.com",
     "Verify the CFITO website URL is current — organization may have updated their web presence."),
    ("SunBiz Filing Fees", "https://sunbiz.org",
     "Current fees are $125 LLC / $70 Corp. Verify these are still accurate."),
    ("USPTO Trademark Fees", "https://www.uspto.gov",
     "TEAS Plus: $250/class, TEAS Standard: $350/class. Verify current fee schedule."),
    ("POK Tech Credential", "TBD",
     "The digital credential issuer (POK Tech) is referenced in Lesson 8-2 but URL is not yet defined. Confirm provider and URL."),
    ("Prospera Website", "https://prosperausa.org",
     "Verify current URL. Some sources show prosperausa.org, others show prospera.org — confirm correct domain."),
    ("Stripe Atlas $500 Fee", "https://stripe.com/atlas",
     "Fee was $500 at time of writing. Verify current pricing."),
    ("SBDC URL Consistency", "https://sbdcorlando.com vs https://sbdcorlando.com",
     "Two different SBDC URLs used in the course. Standardize to one URL — recommend sbdcorlando.com for statewide program, sbdcorlando.com for Orlando-specific."),
]

for name, url, note in review_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run_name = p.add_run(f'• {name} — ')
    run_name.bold = True
    run_name.font.size = Pt(10)
    run_url = p.add_run(url)
    run_url.font.color.rgb = BLUE
    run_url.font.size = Pt(10)
    p2 = doc.add_paragraph(f'  {note}')
    p2.paragraph_format.left_indent = Inches(0.4)
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].font.color.rgb = GRAY
    p2.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# ─── SECTION 5: MISSING RESOURCES ───
add_heading(doc, 'Modules Without External Resources', level=1)
add_para(doc,
    'The following lessons currently reference no external links or organizations. '
    'Consider whether adding 1–2 relevant resources would strengthen the lesson.',
    size=10)
doc.add_paragraph()

no_resource_lessons = [
    ("Module 0 — Lesson 0-1", "U.S. Market Entry Overview", "Could reference sbdcorlando.com or incubator.ucf.edu as next-step resources."),
    ("Module 1 — Lesson 1-3", "Opening Your Business Bank Account", "Opportunity: Mercury, Relay, or Chase Business — practical next-step links."),
    ("Module 2 — Lesson 2-2", "Sales Tax, Accounting & Bookkeeping", "Opportunity: Florida Dept. of Revenue (floridarevenue.com) for sales tax registration."),
    ("Module 3 — Lesson 3-2", "Business Insurance", "Opportunity: NetQuote or an insurance comparison tool for business liability quotes."),
    ("Module 4 — Lesson 4-1", "Contractor vs. Employee", "Opportunity: IRS Worker Classification guide — irs.gov/businesses/small-businesses-self-employed/independent-contractor-defined"),
    ("Module 4 — Lesson 4-3", "Payroll & HR Compliance", "Opportunity: Gusto (gusto.com) — most recommended payroll for small businesses with international founders."),
    ("Modules 5 & 6", "Product, Sales, Culture", "These modules are conceptual by design. Adding resources would be optional."),
    ("Module 7 — Lesson 7-3", "Central Florida by the Numbers", "Data sources could be cited: Orlando Economic Partnership, BLS, U.S. Census Bureau."),
]

for mod, title, note in no_resource_lessons:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{mod} — ')
    run1.bold = True
    run1.font.size = Pt(10)
    run2 = p.add_run(title)
    run2.font.size = Pt(10)
    p2 = doc.add_paragraph(f'  {note}')
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].font.color.rgb = GRAY
    p2.paragraph_format.space_after = Pt(4)

# ─── SAVE ───
out_path = os.path.join(BASE, 'UCF-Resource-Guide.docx')
doc.save(out_path)
print(f'SAVED: {out_path}')
print(f'Total resources documented: {sum(len(v) for v in RESOURCES_BY_LESSON.values())}')
