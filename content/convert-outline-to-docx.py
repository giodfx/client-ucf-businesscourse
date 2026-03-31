from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# --- Style helpers ---
DARK = (0x1A, 0x1A, 0x2E)
BLUE = (0x1A, 0x4E, 0x8C)
MED_BLUE = (0x2E, 0x75, 0xB6)
GRAY = (0x44, 0x44, 0x44)
LIGHT_GRAY = (0x66, 0x66, 0x66)
WHITE = (0xFF, 0xFF, 0xFF)

def sf(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(*DARK)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(*BLUE)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(*MED_BLUE)
    return p

def body(text="", bold=False, italic=False, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        sf(r, bold=bold, italic=italic, color=color)
    return p

def body_rich(parts, space_after=6):
    """parts = list of (text, bold, italic, color)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, color in parts:
        r = p.add_run(text)
        sf(r, bold=bold, italic=italic, color=color)
    return p

def bullet(text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        sf(r1, bold=True)
        r2 = p.add_run(text)
        sf(r2)
    else:
        r = p.add_run(text)
        sf(r)
    return p

def callout(text, bg_color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    sf(r, italic=True, color=LIGHT_GRAY)
    return p

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def outcome_box(text):
    """Outcome line styled distinctly"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run("Outcome: ")
    sf(r1, bold=True, color=BLUE)
    r2 = p.add_run(text)
    sf(r2, italic=True, color=GRAY)


# ============================================================
# TITLE PAGE
# ============================================================

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("U.S. Market Readiness Program")
sf(r, size=26, bold=True, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Course Outline v1")
sf(r, size=16, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run("UCF Business Incubation Program\nInternational Founder Pathway  \u2014  Phase 1")
sf(r, size=12, italic=True, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("Prepared by: Design X Factor\nFebruary 2026")
sf(r, size=11, color=LIGHT_GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run("Status: For Brian\u2019s review + stakeholder validation")
sf(r, size=11, bold=True, color=MED_BLUE)

doc.add_page_break()


# ============================================================
# HOW TO USE THIS DOCUMENT
# ============================================================

heading("How to Use This Document", level=1)

body("This is the shareable outline \u2014 short enough to hand to international entrepreneurs and colleagues for quick feedback. A companion document with the full proposed content will follow once this structure is validated.")

body("We need feedback on:", bold=True, space_after=3)
bullet("Is the story in the right order?")
bullet("Is anything critical missing?")
bullet("Is anything here that doesn\u2019t belong in Phase 1?")
bullet("For entrepreneurs who\u2019ve been through this: What do you wish you\u2019d known?")

hr()

# ============================================================
# PROGRAM OVERVIEW
# ============================================================

heading("Program Overview", level=1)

body("A self-paced program that prepares international entrepreneurs to successfully launch and operate a business in the United States, with a focus on the Central Florida ecosystem.")

body_rich([
    ("This is not a traditional course.", True, False, DARK),
    (" It\u2019s the entry point to a structured pathway:", False, False, None),
])

bullet("Education + readiness \u2014 learn what you need to know", bold_prefix="Phase 1 (this program): ")
bullet("Ecosystem introduction \u2014 warm intros, guided navigation, advisory", bold_prefix="Phase 2 (future): ")
bullet("Full incubator membership \u2014 coaching, mentorship, execution support", bold_prefix="Phase 3 (future): ")

body("")
body_rich([
    ("Target audience: ", True, False, DARK),
    ("International founders vetted by UCF BIP \u2014 established businesses expanding to the U.S., primarily from Latin America, Europe, and Asia.", False, False, None),
])

body_rich([
    ("Estimated time commitment: ", True, False, DARK),
    ("2\u20133 hours total, self-paced", False, False, None),
])

hr()

# ============================================================
# MODULE 1
# ============================================================

heading("Module 1 \u2014 Establishing Your U.S. Business", level=1)
callout("The legal and operational foundation for setting up shop in the United States.")

body("Everything an international founder needs to know to go from \u201cI want to do business here\u201d to having a real, operational U.S. entity.")

body("Topics covered:", bold=True, space_after=3)
bullet("Choosing the right business entity (LLC, C-Corp, S-Corp) \u2014 and why it matters")
bullet("Registering your business in Florida (Sunbiz)")
bullet("Obtaining your EIN (Employer Identification Number)")
bullet("Opening a U.S. business bank account as a foreign national")
bullet("Initial compliance requirements")
bullet("Insurance expectations \u2014 what U.S. partners and clients will expect")
bullet("Contract risk \u2014 how U.S. contract law differs from what you know")
bullet("Regulatory exposure \u2014 industry-specific requirements you may not anticipate")
bullet("Operational liability \u2014 protecting yourself personally")

outcome_box("A clear legal U.S. presence roadmap \u2014 the step-by-step sequence to get legally established.")

hr()

# ============================================================
# MODULE 2
# ============================================================

heading("Module 2 \u2014 How U.S. Business Actually Works", level=1)
callout("The unwritten rules that no one tells you \u2014 until you get it wrong.")

body("The cultural and operational norms that catch international founders off guard. This isn\u2019t theory \u2014 it\u2019s what you need to know to be taken seriously in U.S. business interactions.")

body("Topics covered:", bold=True, space_after=3)
bullet("Speed of decision-making \u2014 why Americans move fast and expect you to")
bullet("Direct communication norms \u2014 what \u201chonest feedback\u201d looks like here")
bullet("Sales expectations \u2014 how U.S. buyers evaluate vendors and partners")
bullet("Trust and credibility signals \u2014 what builds (and destroys) your reputation")
bullet("Meeting etiquette and relationship-building \u2014 it\u2019s different here")
bullet("Founder vs. manager expectations \u2014 what the market expects of you")
bullet("Negotiation style differences \u2014 directness, timelines, follow-through")
bullet("Common cultural missteps \u2014 real examples from international founders")

outcome_box("A cultural readiness checklist \u2014 concrete adjustments to make before your first U.S. business meeting.")

hr()

# ============================================================
# MODULE 3
# ============================================================

heading("Module 3 \u2014 Protecting Your Business", level=1)
callout("Intellectual property, contracts, and knowing when to bring in a professional.")

body("The U.S. has strong IP protections \u2014 but only if you use them. This module covers what international founders commonly miss and what can cost you everything if you ignore it.")

body("Topics covered:", bold=True, space_after=3)
bullet("Overview of U.S. intellectual property: copyright, trademark, patent")
bullet("Common IP mistakes international founders make")
bullet("When and how to bring in an attorney \u2014 and what that costs")
bullet("Protecting your business idea during early conversations")

outcome_box("An IP risk awareness plan \u2014 know what to protect, how, and when to get help.")

hr()

# ============================================================
# MODULE 4
# ============================================================

heading("Module 4 \u2014 Taxes & Financial Fundamentals", level=1)
callout("What you need to know about money in the U.S. \u2014 before you learn it the hard way.")

body("U.S. tax and accounting practices are different from almost everywhere else. This module gives you the baseline financial literacy to avoid expensive surprises.")

body("Topics covered:", bold=True, space_after=3)
bullet("Federal vs. state taxes \u2014 yes, you deal with both")
bullet("Sales tax realities \u2014 it varies by state, county, and product")
bullet("Accounting expectations \u2014 why U.S. bookkeeping differs from global norms")
bullet("Working with a U.S. accountant \u2014 what to expect and what it costs")
bullet("Financial compliance basics for foreign-owned entities")

outcome_box("Basic understanding of U.S.-specific financial obligations so you can have an informed conversation with an accountant.")

hr()

# ============================================================
# MODULE 5
# ============================================================

heading("Module 5 \u2014 Going to Market in the U.S.", level=1)
callout("Why your product probably won\u2019t transfer directly \u2014 and what to do about it.")

body("What works in your home market may not work here. This module covers the realities of U.S. customer acquisition and market entry strategy.")

body("Topics covered:", bold=True, space_after=3)
bullet("Why products and services rarely transfer directly to the U.S. market")
bullet("U.S. customer acquisition expectations \u2014 sales cycles, pricing, trust")
bullet("Adapting your value proposition for American buyers")
bullet("Location and market strategy \u2014 why Central Florida, industry clusters")

outcome_box("A realistic assessment of what needs to change in your product/service and go-to-market approach for the U.S.")

hr()

# ============================================================
# MODULE 6
# ============================================================

heading("Module 6 \u2014 Hiring & Immigration Basics", level=1)
callout("Your first U.S. hire, visa pathways, and what you can\u2019t figure out from Google.")

body("If you\u2019re planning to operate in the U.S., you\u2019ll eventually need people. This module covers the basics of hiring, the contractor-vs-employee distinction that trips up every foreign founder, and an honest overview of visa pathways.")

body("Topics covered:", bold=True, space_after=3)
bullet("Hiring your first U.S. employee \u2014 what\u2019s involved")
bullet("Contractor vs. employee \u2014 the legal distinction that matters enormously")
bullet("Visa pathway overview \u2014 O-1, E-2, L-1, and what\u2019s realistic")
bullet("Payroll expectations and employer obligations")

outcome_box("Enough knowledge to plan your first hire and understand your immigration options. Sets the stage for deeper Phase 2 advisory.")

hr()

# ============================================================
# MODULE 7
# ============================================================

heading("Module 7 \u2014 The Central Florida Ecosystem", level=1)
callout("You don\u2019t have to do this alone \u2014 here\u2019s who can help.")

body("An awareness-level introduction to the organizations, programs, and resources available to international entrepreneurs in Central Florida. This is not a how-to guide \u2014 it\u2019s a menu of what exists so you know what to ask about.")

body("Topics covered:", bold=True, space_after=3)
bullet("Key organizations: SBDC, SBA, chambers of commerce")
bullet("Economic development agencies (Orlando Economic Partnership, SelectFlorida)")
bullet("University partnerships \u2014 how schools like UCF support businesses")
bullet("Industry-specific resources and trade organizations (e.g., CFITO for LATAM trade)")
bullet("UCF Business Incubation Program \u2014 where this fits in the ecosystem")
bullet("Hispanic Chamber, Prospera, and other community-specific resources (where applicable)")

outcome_box("A list of support organizations and their specialties \u2014 know who to contact and what they offer.")

hr()

# ============================================================
# WRAP-UP
# ============================================================

heading("Course Wrap-Up \u2014 Reflection & Next Steps", level=1)
callout("Taking stock of what you\u2019ve learned and preparing for what comes next.")

body("This is not a test. It\u2019s a chance to reflect on what was most valuable, what questions remain, and what your next steps should be.")

body("Includes:", bold=True, space_after=3)
bullet("Self-assessment: What did you learn? What\u2019s still unclear?")
bullet("Value perception feedback \u2014 was this helpful? What would you change?")
bullet("Open question submission \u2014 drop questions for Brian to address during office hours")
bullet("Three exit pathways:")
bullet("Exit with resources \u2014 download your checklists and reference guides", bold_prefix="", indent=0.5)
bullet("Learn more about UCF BIP \u2014 explore Phase 2 ecosystem introduction", bold_prefix="", indent=0.5)
bullet("Schedule time with Brian \u2014 book a 1:1 to discuss your specific situation", bold_prefix="", indent=0.5)

hr()

# ============================================================
# REGIONAL SUPPLEMENTS
# ============================================================

heading("Regional Supplements (Pilot Scope: 2 Regions)", level=1)

body("Throughout the core modules, learners may receive brief supplementary content based on their region of origin. This is not branching \u2014 it\u2019s a thin layer of additional relevance on top of the same core content.")

body("Pilot examples:", bold=True, space_after=3)
bullet("Links to Hispanic Chamber, Prospera, CFITO trade delegations, Spanish-language business resources in Central Florida", bold_prefix="Latin America / Spanish-speaking: ")
bullet("TBD (Asia, Europe, or Canada \u2014 based on Brian\u2019s input on which cohort is most likely in the pilot)", bold_prefix="Second region: ")

hr()

# ============================================================
# OUT OF SCOPE
# ============================================================

heading("What This Document Does NOT Cover", level=1)

body("These items are confirmed as out of scope for Phase 1 but noted for future phases:")

bullet("Phase 2 content (360 assessment, warm introductions, ecosystem navigation)")
bullet("Phase 3 content (incubator membership, coaching, mentorship)")
bullet("Dynamic content branching per country (future \u2014 pilot uses predetermined regions)")
bullet("KPI dashboard for funders (separate discussion, Phase 2+)")
bullet("Payment gateway / registration system (pilot uses DXF platform with code-based access)")
bullet("AI-powered content adaptation (the \u201cempowerment\u201d modality \u2014 planned but scoped separately based on token cost estimates for pilot volume)")

hr()

# ============================================================
# STAKEHOLDER QUESTIONS
# ============================================================

heading("Questions for Stakeholder Validation", level=1)

body("If you\u2019re reviewing this outline, we\u2019d appreciate your input on:")

bullet("Is there a topic that every international founder needs to know that we haven\u2019t included?", bold_prefix="What\u2019s missing? ")
bullet("Is anything here that founders already know or can easily Google?", bold_prefix="What\u2019s unnecessary? ")
bullet("For each module \u2014 do you want a high-level overview or a step-by-step guide?", bold_prefix="What\u2019s the right depth? ")
bullet("Does the sequence make sense, or would you rearrange anything?", bold_prefix="What order would you learn this in? ")
bullet("If you\u2019ve gone through this process \u2014 what surprised you, what was painful, what would have saved you time?", bold_prefix="What do you wish you\u2019d known? ")


# ============================================================
# SAVE
# ============================================================

import os
out = os.path.join(os.path.dirname(__file__), "Course-Outline-v1.docx")
doc.save(out)
print(f"Saved: {out}")
