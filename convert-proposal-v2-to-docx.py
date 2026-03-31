from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# --- Style helpers ---
DARK = (0x1A, 0x1A, 0x2E)
BLUE = (0x1A, 0x4E, 0x8C)
MED_BLUE = (0x2E, 0x75, 0xB6)
GRAY = (0x44, 0x44, 0x44)
LIGHT_GRAY = (0x66, 0x66, 0x66)

def sf(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(*DARK)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(*BLUE)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(*MED_BLUE)
    return p

def body(text="", bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        sf(r, bold=bold, italic=italic)
    return p

def body_rich(parts, space_after=6):
    """parts = list of (text, bold, italic, color)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, color in parts:
        r = p.add_run(text)
        sf(r, bold=bold, italic=italic, color=color)
    return p

def bullet(text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        sf(r1, bold=True)
        r2 = p.add_run(text)
        sf(r2)
    else:
        r = p.add_run(text)
        sf(r)
    return p

def callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '6')
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), '2E75B6')
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(*MED_BLUE)

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def meta(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + " ")
    sf(r1, bold=True, size=10)
    r2 = p.add_run(value)
    sf(r2, size=10)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A4E8C')
        shd.set(qn('w:val'), 'clear')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = t.add_row().cells
        bg = 'F2F7FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            row[ci].text = val
            for run in row[ci].paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg)
            shd.set(qn('w:val'), 'clear')
            row[ci]._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ================================================================
# TITLE
# ================================================================
heading("Professional Services Proposal \u2014 Revised", level=1)

p = doc.add_paragraph()
r = p.add_run("U.S. Market Readiness Program for UCF Business Incubation Program")
sf(r, italic=True, size=12, color=GRAY)
p.paragraph_format.space_after = Pt(12)

rule()

meta("Prepared for:", "Brian Bedrick, University of Central Florida \u2014 Business Incubation Program")
meta("Prepared by:", "Gio, Design X Factor \u2014 gio@designxfactor.com")
meta("Date:", "February 2026")
meta("Version:", "2.0 \u2014 Revised scope based on discovery conversations")
meta("Proposal Valid Through:", "March 31, 2026")

rule()
doc.add_paragraph()

# ================================================================
# 1. EXECUTIVE SUMMARY
# ================================================================
heading("1. Executive Summary", level=2)

body("This revised proposal reflects our discovery conversations and Brian's feedback on the scope, goals, and strategic direction of this engagement.")

body("The original proposal included two deliverables: a lead generation platform ($2,000) and a course ($7,500). Since the lead generation system is being handled through a separate internal solution, this proposal focuses exclusively on the course development effort at the agreed $7,500 investment.")

body_rich([
    ("However, based on our discussions, what UCF BIP needs is not a standard course. It is a ", False, False, None),
    ("U.S. Market Readiness Program", True, False, None),
    (" \u2014 an educational experience that also functions as a qualification tool, a data collection system, and the entry point to a longer-term engagement with international entrepreneurs.", False, False, None),
], space_after=6)

body("This proposal reflects that expanded vision while staying within the agreed budget, made possible by leveraging our proprietary development platform to deliver capabilities that would otherwise require significant custom development or expensive third-party tools.")

# ================================================================
# 2. WHAT CHANGED
# ================================================================
heading("2. What Changed", level=2)

body("Based on our conversations and Brian's written feedback, the following shifts have been incorporated:")

add_table(
    ["Original Scope", "Revised Scope"],
    [
        ["Standard self-paced course", "Phase 1 of a 3-phase platform: education + qualification + on-ramp"],
        ["5 content modules", "5 core modules + profiling intake + reflection/routing at exit"],
        ["Simple registration (name/email)", "Business profiling at course start (stage, industry, origin, goals)"],
        ["Pass/fail knowledge checks", "Reinforcement assessments + end-of-course reflection for stakeholder data"],
        ["Dedicated UCF BIP module", "UCF BIP woven into ecosystem; program routing at exit based on profile"],
        ["LearnDash hosting ($29\u201369/mo paid by UCF)", "Hosted on our platform for the pilot \u2014 no monthly cost to UCF"],
        ["Certificate on completion", "3 exit pathways: exit with resources / learn more / schedule time with Brian"],
    ],
    col_widths=[3.0, 3.5]
)

# ================================================================
# 3. WHAT WE ARE BUILDING
# ================================================================
heading("3. What We Are Building", level=2)

heading("The U.S. Market Readiness Program \u2014 Phase 1", level=3)

body("This is the first phase of what Brian has described as a 3-phase platform journey:")

bullet("Online Course (this proposal): Self-paced education + business profiling + qualification routing", bold_prefix="Phase 1 \u2014 ")
bullet("Ecosystem Introduction (future, manual): Warm introductions, guided navigation, one-on-one sessions", bold_prefix="Phase 2 \u2014 ")
bullet("Incubator Membership (existing): Full UCF BIP program participation", bold_prefix="Phase 3 \u2014 ")

body("")
body("We are building Phase 1. But it is designed as an on-ramp \u2014 not a dead end. Every learner who completes Phase 1 exits with a clear next step, and UCF BIP exits with usable data about who they are and what they need.")

# Learner Journey
heading("The Learner Journey", level=3)

body_rich([("INTAKE", True, False, MED_BLUE)], space_after=2)
bullet("Business profile: 5\u20138 questions (stage, industry, country, goals, language)")

body_rich([("CORE COURSE", True, False, MED_BLUE)], space_after=2)
bullet("Module 1: Getting Started in the U.S.")
bullet("Module 2: U.S. Business Culture")
bullet("Module 3: Protecting Your Business (IP)")
bullet("Module 4: Finding Support & Resources")
bullet("Module 5: Taxes & Financial Fundamentals")

body_rich([("REFLECTION & DATA CAPTURE", True, False, MED_BLUE)], space_after=2)
bullet("What are your top priorities?")
bullet("What questions remain?")
bullet("What are your planned next steps?")

body_rich([("EXIT ROUTING", True, False, MED_BLUE), (" (based on profile + reflection)", False, False, GRAY)], space_after=2)
bullet("Exit \u2014 Resources provided, not a fit for BIP right now", bold_prefix="Path A: ")
bullet("Learn More \u2014 Route to UCF BIP program overview", bold_prefix="Path B: ")
bullet("Connect \u2014 Schedule 15/30/60 min with Brian", bold_prefix="Path C: ")

# Content Modules
heading("Content Modules", level=3)

body("The following modules reflect the agreed outline plus the highest-priority addition from Brian's feedback (Taxes & Financial Fundamentals). Additional modules Brian proposed are positioned as future sprint modules.")

for title, items in [
    ("Module 1 \u2014 Getting Started in the U.S.", [
        "Choosing the right business entity (LLC, C-Corp, S-Corp)",
        "Registering in Florida (Sunbiz), obtaining an EIN",
        "Opening a U.S. business bank account",
        "Initial compliance and insurance requirements",
    ]),
    ("Module 2 \u2014 U.S. Business Culture", [
        "Communication norms and expectations",
        "Meeting etiquette and relationship-building",
        "Negotiation styles and decision-making speed",
        "Common cultural missteps international founders make",
    ]),
    ("Module 3 \u2014 Protecting Your Business (IP)", [
        "Overview of U.S. intellectual property: copyright, trademark, patent",
        "Common mistakes international founders make with IP",
        "When and how to bring in an attorney",
        "Contract basics and risk awareness",
    ]),
    ("Module 4 \u2014 Finding Support & Resources", [
        "Key organizations: SBDC, SBA, chambers of commerce, economic development agencies",
        "What each offers and when to engage them",
        "Navigating the Central Florida ecosystem",
        "UCF BIP as part of the resource landscape",
    ]),
    ("Module 5 \u2014 Taxes & Financial Fundamentals", [
        "Federal and Florida tax obligations for foreign-owned businesses",
        "Sales tax, payroll basics, estimated payments",
        "Financial reporting expectations",
        "When to engage an accountant vs. attorney",
    ]),
]:
    p = doc.add_paragraph()
    r = p.add_run(title)
    sf(r, bold=True)
    p.paragraph_format.space_after = Pt(3)
    for item in items:
        bullet(item)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

body("Content will draw from Brian's provided source materials (OpenMyFloridaBusiness.gov, SBDC, SelectFlorida, CFDC, Orlando Economic Partnership, and others). Principles go in-course; regulatory specifics link out to authoritative external sources.", italic=True)

# ================================================================
# 4. PLATFORM & HOSTING
# ================================================================
heading("4. Platform & Hosting \u2014 Pilot Approach", level=2)

heading("Why Not LearnDash", level=3)
body("The original proposal recommended LearnDash ($29\u201369/month paid by UCF). Based on what we now know about the requirements \u2014 business profiling, routing logic, stakeholder analytics, data export \u2014 a standard LMS like LearnDash cannot support these features without significant custom plugin development.")

heading("What We Are Offering Instead", level=3)
body("For this pilot, we will host the program on our proprietary development platform. This platform includes capabilities that would otherwise require custom development or enterprise-grade LMS tools:")

add_table(
    ["Capability", "What It Does"],
    [
        ["Business Profile Intake", "Configurable questionnaire at course start \u2014 captures stage, industry, origin, goals, language"],
        ["Program Matching", "Rule-based routing that maps profiles to recommended next steps"],
        ["Consultation Booking", "Tiered call-to-action at completion \u2014 integrates with Calendly, Cal.com, or custom link"],
        ["Adaptive Content", "Emphasize or surface content based on learner profile responses"],
        ["Impact Dashboard", "Stakeholder-facing analytics: aggregate profiles, completion rates, satisfaction, referrals"],
        ["Data Export", "Learner data export via CSV or webhook (CRM, spreadsheet, Zapier)"],
        ["Localization", "Multi-language architecture \u2014 Spanish enabled by uploading translated content"],
        ["Payment Processing", "Built-in Stripe and PayPal for collecting the ~$100 course fee"],
        ["Certificates", "Completion certificates with optional blockchain verification"],
    ],
    col_widths=[2.0, 4.5]
)

heading("Pilot Terms", level=3)
bullet("We host, manage, and maintain the platform at no additional cost to UCF BIP during the pilot period.", bold_prefix="For the pilot: ")
bullet("We evaluate together what the best long-term path looks like \u2014 continued hosting on our platform (at a service fee based on usage), migration to a UCF-managed LMS, or SCORM export.", bold_prefix="After the pilot: ")
bullet("All learner data collected belongs to UCF BIP. Full export available at any time.", bold_prefix="Data ownership: ")

heading("What We Need: Expected User Volume", level=3)
body("To configure the platform and set realistic analytics baselines, we need to understand:")
bullet("How many learners do you expect in the first 3 months? (50? 100? 500?)")
bullet("How many in the first year?")
bullet("Cohort-based enrollment (tied to events like Hungarian Hub, April 17) or continuous open enrollment?")
bullet("Target number for stakeholder reporting? (e.g., \"100 international entrepreneurs served in Q2\")")

callout("This does not change the cost \u2014 it helps us configure capacity and design the intake flow for the expected volume.")

# ================================================================
# 5. PROCESS
# ================================================================
heading("5. Our Process", level=2)

add_table(
    ["Stage", "What Happens", "Your Role", "Week"],
    [
        ["Blueprint", "Course architecture, profile questions, routing rules", "Review and approve", "1\u20132"],
        ["Experience Design", "Visual identity, tone, learner journey UX", "Confirm direction", "2\u20133"],
        ["Content Generation", "Lesson content, scenarios, assessments, reflections", "Review; provide feedback", "3\u20134"],
        ["Media Production", "Images, interactives, video segments", "Review samples", "4\u20135"],
        ["Assembly & QA", "Full build, profiling/routing config, accessibility", "Final review", "5\u20136"],
        ["User Validation", "Test with 3\u20135 UCF BIP international companies", "Facilitate access", "6\u20137"],
        ["Launch", "Production deploy, payments, admin orientation", "Confirm readiness", "7\u20138"],
    ],
    col_widths=[1.5, 2.5, 1.5, 0.7]
)

heading("Key Milestones", level=3)
add_table(
    ["Milestone", "Target", "What Brian Can Show"],
    [
        ["Blueprint approved", "Week 2", "Course architecture and profiling system design"],
        ["First module live", "Week 4\u20135", "Working demo for internal review or early validation"],
        ["Hungarian Hub event", "April 17", "Live course available for sign-ups (if timeline holds)"],
        ["Full launch", "Week 8", "Complete program open for enrollment"],
    ],
    col_widths=[1.5, 1.0, 4.0]
)

# ================================================================
# 6. INVESTMENT
# ================================================================
heading("6. Investment", level=2)

heading("Course Development \u2014 $7,500", level=3)

add_table(
    ["Line Item", "Amount"],
    [
        ["Course Design & Architecture (blueprint, experience design, profile intake, routing logic)", "$1,200"],
        ["Content Production (5 modules, assessments, reflection module, exit pathway content)", "$3,200"],
        ["Media Production (images, interactive components, video segments)", "$1,200"],
        ["Technical Implementation (platform config, profiling/routing, analytics, SCORM packaging)", "$1,200"],
        ["User Validation & Iteration (3\u20135 users from current program)", "$400"],
        ["Deployment & Launch (production deploy, admin orientation, payment setup)", "$300"],
        ["Total", "$7,500"],
    ],
    col_widths=[5.0, 1.2]
)

heading("What Is Included (That Was Not in v1)", level=3)
body("At no additional cost over the original proposal:")

for item in [
    "Business profiling intake system (5\u20138 configurable questions)",
    "Program matching and exit routing logic (3 pathways)",
    "Stakeholder impact dashboard (aggregate analytics)",
    "Learner data export (CSV/webhook)",
    "Platform hosting for the pilot period (no LearnDash subscription required)",
    "Payment processing integration (Stripe/PayPal for the $100 course fee)",
    "Completion certificates",
    "Spanish-ready architecture (content translation is a separate effort if desired)",
]:
    bullet(item)

body("")
callout("These capabilities are possible because we are using our own platform, which supports these features natively. On a standard LMS, these would require custom development estimated at $8,000\u2013$15,000 in addition to the course content itself.")

heading("Payment Schedule", level=3)
add_table(
    ["Milestone", "Amount", "When"],
    [
        ["Course Development \u2014 Start", "$3,750", "Upon blueprint approval"],
        ["Course Development \u2014 Completion", "$3,750", "Upon course delivery and launch"],
    ],
    col_widths=[2.5, 1.0, 3.0]
)

heading("Third-Party Costs (Paid by UCF BIP)", level=3)
add_table(
    ["Item", "Estimated Cost", "Notes"],
    [
        ["Domain Name (optional)", "~$12/year", "Custom domain (e.g., learn.ucfbip.com)"],
        ["Payment Processing", "2.9% + $0.30/txn", "Stripe/PayPal fees on ~$100 course fee (~$3.20 per enrollment)"],
    ],
    col_widths=[1.8, 1.2, 3.5]
)

callout("No monthly hosting fees during the pilot. Platform hosting is included in the engagement.")

heading("Optional Add-Ons", level=3)
add_table(
    ["Add-On", "Description", "Estimate"],
    [
        ["Full Spanish Production", "Professional translation + Spanish audio narration", "$1,500"],
        ["Additional Sprint Module", "Standalone module (Go-to-Market, Hiring, Location Strategy)", "$1,200/module"],
        ["Blockchain Certification", "LinkedIn-shareable digital badge via Pok.tech", "$800"],
        ["Post-Pilot Platform Hosting", "Continued hosting, maintenance, support", "TBD"],
    ],
    col_widths=[1.8, 3.2, 1.2]
)

heading("Future Sprint Modules (From Brian's Feedback)", level=3)
body("Brian's feedback included four additional module topics. These are excellent candidates for future development:")
add_table(
    ["Module", "Topic"],
    [
        ["Go-to-Market in the U.S.", "Market entry strategy, distribution, pricing, sales approach"],
        ["Hiring & Immigration Basics", "U.S. employment law, visa considerations, payroll setup"],
        ["Location & Market Strategy", "Site selection, regional advantages, cost of living, logistics"],
        ["UCF BIP Deep Dive", "Detailed program content \u2014 could serve as Phase 2 entry"],
    ],
    col_widths=[2.2, 4.3]
)
body("Each module: $1,200. Can be developed independently without modifying the core course.", italic=True)

# ================================================================
# 7. TERMS
# ================================================================
heading("7. Terms & Conditions", level=2)

for label, text in [
    ("Revisions:", "Two (2) rounds included at each review stage. Additional rounds at $75/hour."),
    ("Content & Materials:", "UCF BIP provides subject matter content and source materials. Brian Bedrick is the single point of contact for all feedback and approvals."),
    ("Intellectual Property:", "Upon full payment, UCF BIP receives full ownership of all course content and educational materials. The platform used to host during the pilot remains the property of Design X Factor. Course content is SCORM-exportable and can be migrated at any time. Design X Factor retains portfolio reference rights."),
    ("Data Ownership:", "All learner data collected through the platform belongs to UCF BIP. Full export available at any time in standard formats (CSV, JSON)."),
    ("Platform Hosting:", "Included for the pilot period at no additional cost. Post-pilot terms agreed upon separately. UCF BIP is never locked in \u2014 course content can be exported as SCORM and deployed elsewhere."),
    ("Pilot Pricing:", "Exclusive to this initial engagement. Future work scoped at standard rates."),
    ("Cancellation:", "Either party may terminate with 14 days written notice. Work completed to date invoiced proportionally."),
]:
    body_rich([(label + " ", True, False, None), (text, False, False, None)], space_after=8)

# ================================================================
# 8. NEXT STEPS
# ================================================================
heading("8. Next Steps", level=2)

bullet("How many learners in the first 3\u20136 months?", bold_prefix="1. Confirm expected user volume \u2014 ")
bullet("Sign-off to begin.", bold_prefix="2. Review and approve this proposal \u2014 ")
bullet("Targeting week of March 3.", bold_prefix="3. Kickoff \u2014 ")
bullet("April 17 as target for live course availability.", bold_prefix="4. Hungarian Hub milestone \u2014 ")

body("")
body("We are ready to begin immediately upon approval.", bold=True)

# ================================================================
# 9. ABOUT
# ================================================================
rule()
heading("9. About Design X Factor", level=2)

body("Design X Factor builds AI-powered digital learning solutions using our proprietary platform. We don't just build courses \u2014 we build learning infrastructure that collects data, qualifies participants, and scales beyond the limits of traditional workshops and static content.")

for label, text in [
    ("Platform-first \u2014 ", "Our own technology delivers capabilities that standard LMS tools cannot, without the cost of custom development."),
    ("Speed \u2014 ", "Full course production in weeks, not months."),
    ("Quality by design \u2014 ", "Mandatory quality gates at every stage."),
    ("SCORM portability \u2014 ", "Every deliverable works across any modern LMS. You are never locked in."),
    ("Built for iteration \u2014 ", "Modular architecture grows with your program."),
]:
    bullet(text, bold_prefix=label)

rule()

body("")
p = doc.add_paragraph()
r = p.add_run("This proposal is submitted in confidence and is intended solely for the use of the University of Central Florida Business Incubation Program.")
sf(r, italic=True, size=9, color=LIGHT_GRAY)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run("Design X Factor\n")
sf(r1, bold=True)
r2 = p.add_run("gio@designxfactor.com\ndesignxfactor.com")
sf(r2, color=GRAY)

output = r"G:\z-CUSTOM_DEV\CourseFuture\clients\UCF\BusinessCourse\UCF-BIP-Proposal-v2.docx"
doc.save(output)
print(f"Saved: {output}")
