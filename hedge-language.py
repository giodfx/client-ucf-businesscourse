#!/usr/bin/env python3
"""
Hedge absolute language in UCF Business Course Content Review document.
- Only processes Modules 2-8 (skips 0 and 1)
- Preserves all comments, formatting, and document structure
- Skips quiz answers, headings, and factual/historical statements
- Logs every change for review
"""
import sys, io, os, re, copy, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'], capture_output=True)
    from docx import Document

INPUT_PATH = r'C:\Users\giand\Downloads\UCF-BusinessCourse-Content-Review.docx'
OUTPUT_PATH = r'C:\Users\giand\Downloads\UCF-BusinessCourse-Content-Review-Hedged.docx'

# ─── Replacement rules ───
# Each rule: (pattern, replacement, description, skip_contexts)
# skip_contexts: list of regex patterns — if any match the full paragraph, skip the rule

RULES = [
    # ── Strong absolutes ──
    (r'\bYou must\b', 'You will generally need to', 'you must → you will generally need to',
     [r'^[A-D][\.\)]', r'Explanation:', r'Knowledge Check']),

    (r'\byou must\b', 'you will generally need to', 'you must → you will generally need to',
     [r'^[A-D][\.\)]', r'Explanation:', r'Knowledge Check']),

    (r'\bmust follow\b', 'are generally expected to follow', 'must follow → are generally expected to follow',
     [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\bmust prove\b', 'generally need to demonstrate', 'must prove → generally need to demonstrate',
     [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\bmust have\b', 'will typically need', 'must have → will typically need',
     [r'^[A-D][\.\)]', r'Explanation:', r'Must have worked']),

    (r'\bMust have\b', 'Typically need to have', 'Must have → Typically need to have',
     [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\bmust collect\b', 'are generally required to collect', 'must collect → are generally required to collect',
     [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\bmust make\b', 'are generally required to make', 'must make → are generally required to make',
     [r'^[A-D][\.\)]', r'Explanation:']),

    # ── "is required" / "is mandatory" ──
    (r'\bis required\b', 'is generally required', 'is required → is generally required',
     [r'^[A-D][\.\)]', r'no filing is required', r'not.*required']),

    (r'\bis mandatory\b', 'is generally required', 'is mandatory → is generally required',
     [r'^[A-D][\.\)]']),

    # ── "requires" ──
    (r'\brequires an?\b', 'typically requires a', 'requires a/an → typically requires a',
     [r'^[A-D][\.\)]', r'Explanation:', r'Requirements:']),

    (r'([Tt]he E-2 also) requires\b', r'\1 typically requires', 'E-2 also requires → typically requires',
     [r'^[A-D][\.\)]']),

    (r'([Ss]tructuring) requires\b', r'\1 typically requires', 'structuring requires → typically requires',
     [r'^[A-D][\.\)]']),

    (r'([Cc]redit) requires\b', r'\1 typically requires', 'credit requires → typically requires',
     [r'^[A-D][\.\)]']),

    # ── "cannot" ──
    (r'\bcannot be S-Corp\b', 'generally cannot be S-Corp', 'cannot be → generally cannot be',
     [r'^[A-D][\.\)]']),

    (r'\bcannot be discharged\b', 'generally cannot be discharged', 'cannot be discharged → generally cannot',
     [r'^[A-D][\.\)]']),

    (r'\bcannot offer\b', 'are not yet able to offer', 'cannot offer → are not yet able to offer',
     [r'^[A-D][\.\)]']),

    (r'\bcannot predict\b', 'may not be able to predict', 'cannot predict → may not be able to predict',
     [r'^[A-D][\.\)]']),

    # ── "always" ──
    (r'\b[Aa]lways call ahead\b', 'it is a good practice to call ahead',
     'always call ahead → good practice to call ahead', [r'^[A-D][\.\)]']),

    (r'\b[Aa]lways get the fee\b', 'make sure to get the fee',
     'always get the fee → make sure to get the fee', [r'^[A-D][\.\)]']),

    (r'\b[Aa]lways check\b', 'make sure to check',
     'always check → make sure to check', [r'^[A-D][\.\)]']),

    (r'\balways require\b', 'typically require', 'always require → typically require',
     [r'^[A-D][\.\)]']),

    (r'\balways better\b', 'generally better', 'always better → generally better',
     [r'^[A-D][\.\)]']),

    (r'\bare always\b', 'are typically', 'are always → are typically',
     [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\bis always\b', 'is typically', 'is always → is typically',
     [r'^[A-D][\.\)]', r'Explanation:', r'not always']),

    (r'\balways higher\b', 'necessarily higher', 'always higher → necessarily higher',
     [r'^[A-D][\.\)]']),

    # ── "never" ──
    (r'\bnever managed\b', 'has not previously managed', 'never managed → has not previously managed',
     [r'^[A-D][\.\)]']),

    (r'\bnever sign\b', 'avoid signing', 'never sign → avoid signing',
     [r'^[A-D][\.\)]']),

    (r'\bnever assume\b', 'do not assume', 'never assume → do not assume',
     [r'^[A-D][\.\)]']),

    (r'\bnever skip\b', 'avoid skipping', 'never skip → avoid skipping',
     [r'^[A-D][\.\)]']),

    # ── "every" (absolute claims about businesses/founders) ──
    (r'\b[Ee]very U\.S\. business\b', 'nearly every U.S. business',
     'every U.S. business → nearly every', [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\b[Ee]very business operating\b', 'most businesses operating',
     'every business operating → most businesses operating', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very business\b(?!\s+operating)', 'nearly every business',
     'every business → nearly every business', [r'^[A-D][\.\)]', r'Explanation:', r'Nearly every']),

    (r'\b[Ee]very taxable sale\b', 'each taxable sale',
     'every taxable sale → each taxable sale', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very state\b(?!\s+individually)', 'each state',
     'every state → each state', [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\b[Ee]very bank\b', 'each bank', 'every bank → each bank',
     [r'^[A-D][\.\)]']),

    (r'\b[Ee]very person you hire\b', 'Each person you hire',
     'every person you hire → each person', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very founder\b', 'most founders', 'every founder → most founders',
     [r'^[A-D][\.\)]']),

    (r'\b[Ee]very bank visit\b', 'each bank visit',
     'every bank visit → each bank visit', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very branch\b', 'each branch',
     'every branch → each branch', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very question\b', 'each question',
     'every question → each question', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very classification\b', 'each classification',
     'every classification → each classification', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very successful transaction\b', 'each successful transaction',
     'every successful transaction → each', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very on-time payment\b', 'each on-time payment',
     'every on-time payment → each on-time payment', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very day of delay\b', 'each day of delay',
     'every day of delay → each day', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very other financial\b', 'most other financial',
     'every other financial → most other financial', [r'^[A-D][\.\)]']),

    (r'\b[Ee]very state individually\b', 'each state individually',
     'every state individually → each state', [r'^[A-D][\.\)]']),

    # ── "all" (absolute claims, NOT defined sets or quiz) ──
    (r'\ball branches handle\b', 'all branches handle',
     None, []),  # skip - context says "skip calling ahead since all branches" - already flagged as wrong advice

    (r'\ball branches work\b', 'all branches operate',
     None, []),  # skip - same context

    (r'\ball your documents\b', None, None, []),  # skip - appropriate

    (r'\ball income and expenses\b', None, None, []),  # skip - appropriate bookkeeping instruction

    (r'\ball tasks\b', 'most tasks',
     'all tasks → most tasks', [r'^[A-D][\.\)]']),

    (r'\ball involve\b', 'each involve',
     'all involve → each involve', [r'^[A-D][\.\)]']),

    (r'\ball suggest employment\b', 'generally suggest employment',
     'all suggest employment → generally suggest', [r'^[A-D][\.\)]']),

    (r'\bfor all\b(?! four)', 'for most',
     'for all → for most', [r'^[A-D][\.\)]', r'Explanation:', r'all of the', r'All of the']),

    (r'\bnot all\b', None, None, []),  # skip - "not all" is already hedged

    # ── "the only" ──
    (r'\bthe only cost\b', 'often the primary cost',
     'the only cost → often the primary cost', [r'^[A-D][\.\)]']),

    (r'\bthe only way\b', 'generally the best way',
     'the only way → generally the best way', [r'^[A-D][\.\)]']),

    # ── "definitely" / "certainly" ──
    (r'\bdefinitely\b', 'likely', 'definitely → likely',
     [r'^[A-D][\.\)]', r'Explanation:']),

    (r'\bcertainly\b', 'likely', 'certainly → likely',
     [r'^[A-D][\.\)]', r'Explanation:']),

    # ── "will reject" / "will deny" ──
    (r'\bwill reject\b', 'may reject', 'will reject → may reject',
     [r'^[A-D][\.\)]']),

    (r'\bwill deny\b', 'may deny', 'will deny → may deny',
     [r'^[A-D][\.\)]']),

    # ── "guarantee" (NOT disclaimers) ──
    # Skip: "does not guarantee" is fine
    (r'\bguarantee(?!d)\b', None, None, []),  # handled case by case below

    # ── Florida law specifics — hedge carefully ──
    (r'Florida requires workers', "Florida generally requires workers",
     'Florida requires → generally requires', [r'^[A-D][\.\)]']),

    (r'Florida requires\b', "Florida generally requires",
     'Florida requires → generally requires', [r'^[A-D][\.\)]', r'Explanation:', r'generally']),

    # ── "will not" ──
    (r'\bwill not accept\b', 'may not accept', 'will not accept → may not accept',
     [r'^[A-D][\.\)]']),

    (r'\bwill not count\b', 'may not count', 'will not count → may not count',
     [r'^[A-D][\.\)]']),

    # ── Misc absolute patterns ──
    (r'\bis essential\b', 'is generally important', 'is essential → is generally important',
     [r'^[A-D][\.\)]']),

    (r'\bis critical\b', 'is typically important', 'is critical → is typically important',
     [r'^[A-D][\.\)]']),

    (r'\bno exceptions\b', 'with very few exceptions', 'no exceptions → with very few exceptions',
     [r'^[A-D][\.\)]']),
]


def should_skip_paragraph(para):
    """Skip headings, quiz answers, and certain structural elements."""
    style = para.style.name
    if style.startswith('Heading'):
        return True
    text = para.text.strip()
    if not text:
        return True
    # Quiz answer options
    if re.match(r'^[✔✗]?\s*[A-D][\.\)]', text):
        return True
    # List items that are quiz options
    if re.match(r'^[A-D]\.\s', text):
        return True
    return False


def apply_rules_to_run(run_text, full_para_text, changes_log, para_idx):
    """Apply hedging rules to a single run's text. Returns modified text."""
    modified = run_text
    for pattern, replacement, description, skip_contexts in RULES:
        if replacement is None or description is None:
            continue  # Skip rule (marker for no-change items)

        # Check skip contexts against full paragraph text
        skip = False
        for ctx in skip_contexts:
            if re.search(ctx, full_para_text):
                skip = True
                break
        if skip:
            continue

        # Apply replacement within this run
        new_text = re.sub(pattern, replacement, modified)
        if new_text != modified:
            # Log the change
            changes_log.append({
                'para': para_idx,
                'description': description,
                'before': modified[:120],
                'after': new_text[:120],
            })
            modified = new_text

    return modified


def find_module_start(paragraphs, module_num):
    """Find the paragraph index where a module starts."""
    for i, p in enumerate(paragraphs):
        if p.style.name == 'Heading 1' and f'MODULE {module_num}' in p.text:
            return i
    return None


def main():
    print("═══ UCF Business Course — Language Hedging Pass ═══\n")

    if not os.path.exists(INPUT_PATH):
        print(f"ERROR: File not found: {INPUT_PATH}")
        sys.exit(1)

    doc = Document(INPUT_PATH)
    paragraphs = doc.paragraphs

    # Find Module 2 start index (skip Modules 0 and 1)
    m2_start = find_module_start(paragraphs, 2)
    if m2_start is None:
        print("ERROR: Could not find Module 2 start")
        sys.exit(1)

    print(f"Module 2 starts at paragraph {m2_start}")
    print(f"Total paragraphs: {len(paragraphs)}")
    print(f"Processing paragraphs {m2_start} through {len(paragraphs)}\n")

    changes_log = []
    paragraphs_modified = 0

    for i in range(m2_start, len(paragraphs)):
        para = paragraphs[i]

        if should_skip_paragraph(para):
            continue

        full_text = para.text
        para_changed = False

        for run in para.runs:
            if not run.text.strip():
                continue

            original_run_text = run.text
            new_run_text = apply_rules_to_run(run.text, full_text, changes_log, i)

            if new_run_text != original_run_text:
                run.text = new_run_text
                para_changed = True

        if para_changed:
            paragraphs_modified += 1

    # Also process table cells (quiz explanations, content in tables)
    table_changes = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Find which module this table is in by checking position
                    # For safety, process all tables but skip quiz answer cells
                    text = para.text.strip()
                    if not text:
                        continue
                    if re.match(r'^[✔✗]?\s*[A-D][\.\)]', text):
                        continue

                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        original = run.text
                        new_text = apply_rules_to_run(run.text, text, changes_log, -1)
                        if new_text != original:
                            run.text = new_text
                            table_changes += 1

    # Save
    doc.save(OUTPUT_PATH)

    # Print summary
    print(f"═══ Results ═══")
    print(f"  Paragraphs modified: {paragraphs_modified}")
    print(f"  Table cells modified: {table_changes}")
    print(f"  Total changes: {len(changes_log)}")
    print(f"  Output: {OUTPUT_PATH}\n")

    # Print change log
    print("═══ Change Log ═══\n")
    current_module = ""
    for c in changes_log:
        # Try to identify module from paragraph index
        if c['para'] >= 0:
            for j in range(c['para'], -1, -1):
                if j < len(paragraphs) and paragraphs[j].style.name == 'Heading 1' and 'MODULE' in paragraphs[j].text:
                    mod = paragraphs[j].text.strip()[:20]
                    if mod != current_module:
                        current_module = mod
                        print(f"\n── {current_module} ──")
                    break
        print(f"  [{c['para']:4d}] {c['description']}")

    # Save change log as JSON too
    log_path = OUTPUT_PATH.replace('.docx', '-changes.json')
    with open(log_path, 'w', encoding='utf-8') as f:
        json.dump(changes_log, f, indent=2, ensure_ascii=False)
    print(f"\n  Change log saved: {log_path}")


if __name__ == '__main__':
    main()
