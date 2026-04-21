#!/usr/bin/env python3
"""
Generate a Word document for Brian Bedrick's content review.
Reads all module JSON files and renders every content block into a
structured, readable .docx with:
  - Cover page with module/lesson table of contents
  - Each module on a new page
  - All block types rendered appropriately
"""
import json, os, sys, io, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from bs4 import BeautifulSoup, NavigableString

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
BASE = os.path.dirname(os.path.abspath(__file__))

# ── Module files in order ──
MODULE_FILES = [
    ('Module 0', 'generated/content-phase1-module0.json'),
    ('Module 1', 'generated/content-phase1-module1.json'),
    ('Module 2', 'generated/content-phase1-module2.json'),
    ('Module 3', 'generated/content-phase1-module3.json'),
    ('Module 4', 'generated/content-phase1-module4.json'),
    ('Module 5', 'generated/content-phase1-module5.json'),
    ('Module 6', 'generated/content-phase1-module6.json'),
    ('Module 7', 'generated/content-phase1-module7.json'),
    ('Module 8', 'generated/content-phase1-module8.json'),
]

# Callout type labels
CALLOUT_LABELS = {
    'fun-fact': 'Fun Fact',
    'funfact': 'Fun Fact',
    'common-mistakes': 'Common Mistakes',
    'cm': 'Common Mistakes',
    'watch-money': 'Watch Your Money',
    'wm': 'Watch Your Money',
    'tip': 'Tip',
    'warning': 'Warning',
    'disclaimer': 'Disclaimer',
    'bys': 'Before You Start',
}

# ── MODULE LOCATION NAMES (from April 6 navigation decision) ──
MODULE_LOCATIONS = {
    0: 'Highway',
    1: 'Downtown',
    2: 'Financial District',
    3: 'Space Center',
    4: 'Theme Park District',
    5: 'Harbor',
    6: 'Beach',
    7: 'Natural Springs',
    8: 'UCF Campus',
}

MODULE_DESCRIPTIONS_SHORT = {
    0: 'Get clear on what it actually takes to enter the U.S. market',
    1: 'Learn how to structure, register, and open your U.S. business',
    2: 'Learn how the U.S. tax and banking system works for your business',
    3: 'Shield your business with the right IP, insurance, and contracts',
    4: 'Build your U.S. team without costly legal and compliance mistakes',
    5: 'Learn how to find, reach, and close your first U.S. customers',
    6: 'Discover how Americans communicate, negotiate, and make decisions',
    7: 'Connect with the free resources most international founders overlook',
    8: 'Assess your readiness and plan your next steps with UCF BIP',
}

# ── COLORS ──
BLUE = RGBColor(0x1A, 0x3A, 0x5C)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0xC0, 0x6D, 0x2B)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED_ACCENT = RGBColor(0xB7, 0x1C, 0x1C)


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    # Heading 1 = Module title
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2 = Lesson title
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(16)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 = Block title / section
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)
    h3.font.color.rgb = BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # Heading 4 = Sub-section within block
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(11)
    h4.font.color.rgb = DARK
    h4.font.bold = True
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)


def add_horizontal_line(doc):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_callout_box(doc, label, block_type_hint=''):
    """Add a colored label paragraph for callout blocks."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'[{label}]')
    run.bold = True
    run.font.size = Pt(10)
    if 'fun' in label.lower():
        run.font.color.rgb = GREEN
    elif 'mistake' in label.lower():
        run.font.color.rgb = RED_ACCENT
    elif 'money' in label.lower() or 'watch' in label.lower():
        run.font.color.rgb = ACCENT
    else:
        run.font.color.rgb = BLUE


def render_html_to_doc(doc, html_content, indent=0):
    """Parse HTML content and render into Word paragraphs."""
    if not html_content:
        return

    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                p = doc.add_paragraph(text)
                if indent:
                    p.paragraph_format.left_indent = Cm(indent)
            continue

        tag = element.name

        if tag in ('p',):
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.left_indent = Cm(indent)
            _render_inline(p, element)

        elif tag in ('h3',):
            p = doc.add_paragraph(element.get_text(), style='Heading 3')

        elif tag in ('h4',):
            p = doc.add_paragraph(element.get_text(), style='Heading 4')

        elif tag in ('h5', 'h6'):
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(11)

        elif tag == 'ul':
            _render_list(doc, element, ordered=False, indent=indent)

        elif tag == 'ol':
            _render_list(doc, element, ordered=True, indent=indent)

        elif tag == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(4)
            run = p.add_run(element.get_text())
            run.italic = True
            run.font.color.rgb = GRAY

        elif tag == 'table':
            _render_html_table(doc, element)

        elif tag in ('strong', 'b', 'em', 'i', 'a', 'span'):
            # Inline element at top level — wrap in paragraph
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.left_indent = Cm(indent)
            _render_inline(p, element)

        elif tag == 'br':
            pass

        elif tag == 'div':
            render_html_to_doc(doc, element.decode_contents(), indent)

        else:
            text = element.get_text().strip()
            if text:
                p = doc.add_paragraph(text)
                if indent:
                    p.paragraph_format.left_indent = Cm(indent)


def _render_inline(paragraph, element):
    """Render inline HTML elements into a Word paragraph."""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                paragraph.add_run(text)
        elif child.name in ('strong', 'b'):
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ('em', 'i'):
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'a':
            text = child.get_text()
            href = child.get('href', '')
            if href:
                run = paragraph.add_run(f'{text} ({href})')
            else:
                run = paragraph.add_run(text)
            run.font.color.rgb = BLUE
            run.underline = True
        elif child.name == 'br':
            paragraph.add_run('\n')
        elif child.name == 'span':
            _render_inline(paragraph, child)
        elif child.name in ('ul', 'ol'):
            pass  # nested lists handled separately
        else:
            text = child.get_text()
            if text:
                paragraph.add_run(text)


def _render_list(doc, list_element, ordered=False, indent=0, level=0):
    """Render ul/ol into Word list items."""
    base_indent = 1.0 + indent + (level * 0.8)
    for i, li in enumerate(list_element.find_all('li', recursive=False)):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(base_indent)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

        prefix = f'{i+1}. ' if ordered else '\u2022 '
        p.add_run(prefix)
        _render_inline(p, li)

        for nested in li.find_all(['ul', 'ol'], recursive=False):
            _render_list(doc, nested, ordered=(nested.name == 'ol'),
                        indent=indent, level=level+1)


def _render_html_table(doc, table_element):
    """Render an HTML table element into a Word table."""
    rows_data = []
    for tr in table_element.find_all('tr'):
        cells = []
        for td in tr.find_all(['td', 'th']):
            cells.append(td.get_text().strip())
        if cells:
            rows_data.append(cells)
    if not rows_data:
        return
    max_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                table.rows[i].cells[j].text = cell_text


def render_markdown_to_doc(doc, md_content, indent=0):
    """Render simple markdown content (used in tabs, etc.)."""
    if not md_content:
        return
    lines = md_content.split('\n')
    current_text = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_text:
                p = doc.add_paragraph()
                if indent:
                    p.paragraph_format.left_indent = Cm(indent)
                text = ' '.join(current_text)
                _render_md_inline(p, text)
                current_text = []
            continue

        if stripped.startswith('**') and stripped.endswith('**'):
            if current_text:
                p = doc.add_paragraph()
                if indent:
                    p.paragraph_format.left_indent = Cm(indent)
                _render_md_inline(p, ' '.join(current_text))
                current_text = []
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.left_indent = Cm(indent)
            run = p.add_run(stripped.strip('*'))
            run.bold = True
        else:
            current_text.append(stripped)

    if current_text:
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        _render_md_inline(p, ' '.join(current_text))


def _render_md_inline(paragraph, text):
    """Handle **bold** and *italic* in markdown text."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ── BLOCK RENDERERS ──────────────────────────────────────────────────────

def render_text_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')
    render_html_to_doc(doc, block.get('content', ''))


def render_callout_block(doc, block):
    block_id = block.get('id', '')
    callout_type = block.get('calloutType', '')

    if not callout_type:
        for suffix, label in CALLOUT_LABELS.items():
            if block_id.endswith(f'-{suffix}') or f'-{suffix}' in block_id:
                callout_type = suffix
                break

    label = CALLOUT_LABELS.get(callout_type, callout_type.replace('-', ' ').title() if callout_type else 'Note')
    title = block.get('title', '')

    add_callout_box(doc, label)
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
    render_html_to_doc(doc, block.get('content', ''))


def render_key_takeaway(doc, block):
    add_callout_box(doc, 'Key Takeaway')
    title = block.get('title', '')
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
    render_html_to_doc(doc, block.get('content', ''))


def render_knowledge_check(doc, block):
    doc.add_paragraph('Knowledge Check', style='Heading 3')
    title = block.get('title', '')
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.italic = True

    for qi, q in enumerate(block.get('questions', []), 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f'Q{qi}: {q.get("stem", "")}')
        run.bold = True

        correct = q.get('correctAnswer', -1)
        for oi, opt in enumerate(q.get('options', [])):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            marker = '\u2714 ' if oi == correct else '   '
            letter = chr(65 + oi)
            run = p.add_run(f'{marker}{letter}. {opt}')
            if oi == correct:
                run.bold = True
                run.font.color.rgb = GREEN

        explanation = q.get('explanation', '')
        if explanation:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run('Explanation: ')
            run.bold = True
            run.font.color.rgb = GRAY
            run2 = p.add_run(explanation)
            run2.font.color.rgb = GRAY
            run2.font.size = Pt(10)


def render_comparison_card(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    columns = block.get('columns', [])
    if not columns:
        return

    all_labels = []
    for col in columns:
        for item in col.get('items', []):
            lbl = item.get('label', '')
            if lbl not in all_labels:
                all_labels.append(lbl)

    num_cols = len(columns) + 1
    num_rows = len(all_labels) + 1

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.rows[0].cells[0].text = ''
    for ci, col in enumerate(columns):
        table.rows[0].cells[ci + 1].text = col.get('heading', '')
        for p in table.rows[0].cells[ci + 1].paragraphs:
            for r in p.runs:
                r.bold = True

    for ri, label in enumerate(all_labels):
        table.rows[ri + 1].cells[0].text = label
        for p in table.rows[ri + 1].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
        for ci, col in enumerate(columns):
            val = ''
            for item in col.get('items', []):
                if item.get('label') == label:
                    val = item.get('value', '')
                    break
            table.rows[ri + 1].cells[ci + 1].text = val

    doc.add_paragraph()


def render_tabs_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    for tab in block.get('tabs', []):
        p = doc.add_paragraph()
        run = p.add_run(f'\u25B6 {tab.get("label", "")}')
        run.bold = True
        run.font.color.rgb = BLUE

        content = tab.get('content', '')
        if '<' in content:
            render_html_to_doc(doc, content, indent=0.5)
        else:
            render_markdown_to_doc(doc, content, indent=0.5)


def render_accordion_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    for section in block.get('sections', []):
        p = doc.add_paragraph()
        run = p.add_run(f'\u25BC {section.get("title", section.get("heading", ""))}')
        run.bold = True

        content = section.get('content', '')
        if '<' in content:
            render_html_to_doc(doc, content, indent=0.5)
        else:
            render_markdown_to_doc(doc, content, indent=0.5)


def render_flashcards_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    for ci, card in enumerate(block.get('cards', []), 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Card {ci} \u2014 ')
        run.bold = True
        run.font.color.rgb = BLUE
        run2 = p.add_run(card.get('front', ''))
        run2.bold = True

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.add_run(card.get('back', ''))


def render_scenario_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    desc = block.get('description', '')
    if desc:
        p = doc.add_paragraph()
        run = p.add_run(desc)
        run.italic = True

    for path in block.get('paths', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f'\u2192 "{path.get("label", "")}"')
        run.bold = True
        run.font.color.rgb = BLUE

        outcome = path.get('outcome', '')
        if outcome:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.add_run(outcome)


def render_table_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    tdata = block.get('table', {})
    headers = tdata.get('headers', [])
    rows = tdata.get('rows', [])
    if not headers and not rows:
        return

    num_cols = max(len(headers), max((len(r) for r in rows), default=0))
    num_rows = len(rows) + (1 if headers else 0)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if headers:
        for ci, h in enumerate(headers):
            if ci < num_cols:
                table.rows[0].cells[ci].text = h
                for p in table.rows[0].cells[ci].paragraphs:
                    for r in p.runs:
                        r.bold = True

    offset = 1 if headers else 0
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            if ci < num_cols:
                table.rows[ri + offset].cells[ci].text = str(cell)

    doc.add_paragraph()


def render_checklist_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    for item in block.get('items', []):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        label = item if isinstance(item, str) else item.get('label', str(item))
        p.add_run(f'\u2610 {label}')


def render_infographic_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    subtitle = block.get('subtitle', '')
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.italic = True
        run.font.color.rgb = GRAY

    for si, step in enumerate(block.get('steps', []), 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        step_title = step.get('title', step.get('label', ''))
        run = p.add_run(f'Step {si}: {step_title}')
        run.bold = True

        content = step.get('content', step.get('description', ''))
        if content:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.add_run(content)

    footer = block.get('footer', '')
    if footer:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(footer)
        run.italic = True
        run.font.size = Pt(10)


def render_timeline_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    for step in block.get('steps', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        label = step.get('label', '')
        step_title = step.get('title', '')
        run = p.add_run(f'{label} \u2014 {step_title}' if label else step_title)
        run.bold = True

        content = step.get('content', '')
        if content:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.add_run(content)


def render_matching_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    instruction = block.get('instruction', '')
    if instruction:
        p = doc.add_paragraph()
        run = p.add_run(instruction)
        run.italic = True

    pairs = block.get('pairs', [])
    if pairs:
        num_rows = len(pairs) + 1
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Phrase'
        table.rows[0].cells[1].text = 'Meaning'
        for p in table.rows[0].cells[0].paragraphs:
            for r in p.runs: r.bold = True
        for p in table.rows[0].cells[1].paragraphs:
            for r in p.runs: r.bold = True
        for i, pair in enumerate(pairs):
            table.rows[i+1].cells[0].text = pair.get('left', '')
            table.rows[i+1].cells[1].text = pair.get('right', '')
        doc.add_paragraph()


def render_data_visualization(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    p = doc.add_paragraph()
    run = p.add_run('[Interactive data visualization \u2014 see online course]')
    run.italic = True
    run.font.color.rgb = GRAY

    html = block.get('htmlContent', '')
    if html:
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text(separator='\n', strip=True)
        if text and len(text) > 20:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.5)
            if len(text) > 1000:
                text = text[:1000] + '...'
            p2.add_run(text)


def render_simulation_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')

    desc = block.get('description', '')
    if desc:
        doc.add_paragraph(desc)

    p = doc.add_paragraph()
    run = p.add_run('[Interactive simulation \u2014 see online course]')
    run.italic = True
    run.font.color.rgb = GRAY


def render_exercise_block(doc, block):
    title = block.get('title', '')
    if title:
        doc.add_paragraph(title, style='Heading 3')
    render_html_to_doc(doc, block.get('content', block.get('instructions', '')))


def render_image_block(doc, block):
    alt = block.get('alt', block.get('title', 'Image'))
    p = doc.add_paragraph()
    run = p.add_run(f'[Image: {alt}]')
    run.italic = True
    run.font.color.rgb = GRAY


# ── MAIN RENDERER ─────────────────────────────────────────────────────────

BLOCK_RENDERERS = {
    'text': render_text_block,
    'callout': render_callout_block,
    'key-takeaway': render_key_takeaway,
    'knowledge-check': render_knowledge_check,
    'comparison-card': render_comparison_card,
    'tabs': render_tabs_block,
    'accordion': render_accordion_block,
    'flashcards': render_flashcards_block,
    'scenario': render_scenario_block,
    'table': render_table_block,
    'checklist': render_checklist_block,
    'infographic': render_infographic_block,
    'timeline': render_timeline_block,
    'matching': render_matching_block,
    'data-visualization': render_data_visualization,
    'simulation': render_simulation_block,
    'exercise': render_exercise_block,
    'image': render_image_block,
    'in-lesson-image': render_image_block,
}


def main():
    doc = Document()
    setup_styles(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER PAGE ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run('UCF Business Incubation Program')
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Starting a Business in the United States')
    run2.font.size = Pt(18)
    run2.font.color.rgb = DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(6)
    run3 = p3.add_run('Content Review Document \u2014 April 2026')
    run3.font.size = Pt(12)
    run3.font.color.rgb = GRAY

    doc.add_paragraph()
    add_horizontal_line(doc)
    doc.add_paragraph()

    # ── TABLE OF CONTENTS ──
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_toc = toc_title.add_run('Course Overview')
    run_toc.font.size = Pt(16)
    run_toc.bold = True
    run_toc.font.color.rgb = BLUE

    doc.add_paragraph()

    # Load all modules and build TOC
    all_modules = []
    for mod_label, fpath in MODULE_FILES:
        full = os.path.join(BASE, fpath)
        if not os.path.exists(full):
            continue
        data = json.load(open(full, encoding='utf-8'))
        mod_title = data.get('moduleTitle', data.get('title', mod_label))
        lessons = data.get('lessons', [data])

        # Append lesson-7-5 (standalone file not in module7.json)
        if mod_label == 'Module 7':
            l75_path = os.path.join(BASE, 'phase1-lessons/lesson-7-5.json')
            if os.path.exists(l75_path):
                l75 = json.load(open(l75_path, encoding='utf-8'))
                lessons.append(l75)

        mod_num = int(mod_label.split()[-1])
        location = MODULE_LOCATIONS.get(mod_num, '')
        short_desc = MODULE_DESCRIPTIONS_SHORT.get(mod_num, '')
        all_modules.append((mod_label, mod_title, lessons, data, mod_num))

        # TOC — module with location name
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f'MODULE {mod_num}  \u2014  {location.upper()}')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = BLUE

        # Module subtitle
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(2)
        run_sub = p_sub.add_run(short_desc)
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = GRAY

        # TOC — lessons
        for lesson in lessons:
            ltitle = lesson.get('title', '')
            p_l = doc.add_paragraph()
            p_l.paragraph_format.left_indent = Cm(1.0)
            p_l.paragraph_format.space_before = Pt(1)
            p_l.paragraph_format.space_after = Pt(1)
            run_l = p_l.add_run(f'\u2022 {ltitle}')
            run_l.font.size = Pt(10)
            run_l.font.color.rgb = DARK

    # ── MODULE CONTENT ──
    for mod_label, mod_title, lessons, data, mod_num in all_modules:
        location = MODULE_LOCATIONS.get(mod_num, '')
        short_desc = MODULE_DESCRIPTIONS_SHORT.get(mod_num, '')

        doc.add_page_break()
        doc.add_paragraph(f'MODULE {mod_num}  \u2014  {location.upper()}', style='Heading 1')

        # Module subtitle line
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(4)
        run_sub = p_sub.add_run(short_desc)
        run_sub.font.size = Pt(12)
        run_sub.font.color.rgb = GRAY
        run_sub.italic = True

        mod_desc = data.get('moduleDescription', data.get('description', ''))
        if mod_desc:
            p = doc.add_paragraph()
            run = p.add_run(mod_desc)
            run.italic = True
            run.font.color.rgb = GRAY

        for lesson in lessons:
            lid = lesson.get('lessonId', '')
            ltitle = lesson.get('title', '')

            add_horizontal_line(doc)
            doc.add_paragraph(ltitle, style='Heading 2')

            # Learning objectives
            objectives = lesson.get('learningObjectives', [])
            if objectives:
                p_obj = doc.add_paragraph()
                run_obj = p_obj.add_run('Learning Objectives')
                run_obj.bold = True
                run_obj.font.color.rgb = GRAY
                run_obj.font.size = Pt(10)
                for obj in objectives:
                    p_o = doc.add_paragraph()
                    p_o.paragraph_format.left_indent = Cm(0.5)
                    p_o.paragraph_format.space_before = Pt(1)
                    p_o.paragraph_format.space_after = Pt(1)
                    obj_text = obj if isinstance(obj, str) else obj.get('text', str(obj))
                    p_o.add_run(f'\u2022 {obj_text}')
                doc.add_paragraph()

            # Content blocks
            for block in lesson.get('contentBlocks', []):
                btype = block.get('type', '')
                renderer = BLOCK_RENDERERS.get(btype)
                if renderer:
                    try:
                        renderer(doc, block)
                    except Exception as e:
                        p = doc.add_paragraph()
                        run = p.add_run(f'[Error rendering {btype} block {block.get("id","")}: {e}]')
                        run.font.color.rgb = RED_ACCENT
                else:
                    content = block.get('content', '')
                    if content:
                        doc.add_paragraph(f'[{btype}]', style='Heading 3')
                        render_html_to_doc(doc, content)

            doc.add_paragraph()

    # Save
    output_path = os.path.join(BASE, 'output', 'UCF-BusinessCourse-Content-Review.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'Document saved: {output_path}')
    print(f'Modules: {len(all_modules)}')
    total_lessons = sum(len(m[2]) for m in all_modules)
    print(f'Lessons: {total_lessons}')


if __name__ == '__main__':
    main()
