from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# --- Style helpers ---
def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x4E, 0x8C)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def body(text="", bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def callout(text):
    """Styled blockquote-style callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '6')
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), '2E75B6')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def meta_line(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + " ")
    set_font(r1, bold=True, size=10)
    r2 = p.add_run(value)
    set_font(r2, size=10)
    return p

# ================================================================
# DOCUMENT CONTENT
# ================================================================

# Title block
heading("UCF BIP — Pre-Production Brief", level=1)

p = doc.add_paragraph()
run = p.add_run('"How to Do Business in the United States" Course')
run.font.italic = True
run.font.size = Pt(12)
run.font.name = "Calibri"
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
p.paragraph_format.space_after = Pt(10)

meta_line("From:", "Design X Factor (Gio)")
meta_line("To:", "Brian Bedrick, UCF Business Incubation Program")
meta_line("Date:", "February 19, 2026")
meta_line("Please review by:", "Monday, February 23, 2026")

rule()
doc.add_paragraph()

p = body("Hi Brian,")
p = doc.add_paragraph()
run = p.add_run(
    "Before our meeting Thursday I need four things from you. "
    "All I'm looking for right now is the structure — not the full content. "
    "Just enough to confirm we're building the right thing."
)
set_font(run)
p.paragraph_format.space_after = Pt(12)

# ================================================================
# SECTION 1
heading("1. Course Outline — Please Review & Confirm", level=2)
body("I've drafted a working outline for both parts of the course. "
     "Take a look and tell me what's off, what's missing, or what should move.", space_after=10)

# Part A
heading('Part A \u2014 "How to Do Business in the United States"', level=3)
p = body("The core business curriculum for international entrepreneurs", italic=True, space_after=8)

p = doc.add_paragraph()
r = p.add_run("Module 1 — Getting Started in the U.S.")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
for item in [
    "Choosing the right business entity (LLC, C-Corp, S-Corp)",
    "Registering your business in Florida (Sunbiz)",
    "Obtaining your EIN",
    "Opening a U.S. business bank account",
    "Initial compliance requirements",
]:
    bullet(item)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Module 2 — U.S. Business Culture")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
for item in [
    "Communication norms and expectations",
    "Meeting etiquette and relationship-building",
    "Negotiation style differences",
    "What surprises most international founders",
]:
    bullet(item)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Module 3 — Protecting Your Business (IP)")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
for item in [
    "Overview of U.S. intellectual property: copyright, trademark, patent",
    "Common mistakes international founders make",
    "When and how to bring in an attorney",
]:
    bullet(item)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Module 4 — Finding Support & Resources")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
for item in [
    "Key organizations: SBDC, SBA, chambers of commerce, economic development agencies",
    "Navigating the Central Florida ecosystem",
]:
    bullet(item)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Module 5 — UCF BIP Programs")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
for item in [
    "Traction, Growth, and Soft Landing — what each is and who it's for",
    "How to get started",
]:
    bullet(item)

doc.add_paragraph()
callout("Does this outline feel right? What would you add, remove, or reorder?")

rule()

# Part B
heading("Part B — UCF BIP Program Content", level=3)
body("The institutional content — UCF's programs and overview. This content comes from you.", italic=True, space_after=8)

p = doc.add_paragraph()
r = p.add_run("UCF BIP Overview")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
bullet("Mission, history, and impact")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("The Programs")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
bullet("brief description, who it's for", bold_prefix="Traction — ")
bullet("brief description, who it's for", bold_prefix="Growth — ")
bullet("brief description, who it's for", bold_prefix="Soft Landing — ")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Next Step for the Learner")
set_font(r, bold=True)
p.paragraph_format.space_after = Pt(3)
bullet("How to connect with UCF BIP after completing the course")

doc.add_paragraph()
callout("Can you fill in a few sentences for each program? We'll flesh it out together — but any notes ahead of time will save us time Thursday.")

# ================================================================
# SECTION 2
rule()
heading("2. Source Materials", level=2)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run("For Part A")
set_font(r1, bold=True)
r2 = p.add_run(
    ", if you have any existing materials you share with international clients — "
    "guides, checklists, links, intake docs, anything — please forward them. "
    "I'd like to review before Thursday so we know what we already have versus what we need to build. "
    "If you don't have anything ready, no problem — I can pull from public references "
    "(SBA, IRS, Florida Division of Corporations, etc.) and we'll fill the gaps together."
)
set_font(r2)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r1 = p.add_run("For Part B")
set_font(r1, bold=True)
r2 = p.add_run(
    ", that content comes entirely from you and we'll build it together during production."
)
set_font(r2)

# ================================================================
# SECTION 3 — Profile Map
rule()
heading("3. Business Profile → UCF BIP Program Map", level=2)

body(
    "This is important for the course experience. The course will guide each learner toward the right "
    "UCF BIP program based on their business profile. I've drafted a starting map below — "
    "please validate it, correct anything that's off, and add profiles I'm missing.",
    space_after=6
)
body(
    "This mapping drives the personalized learning paths in the course, "
    "so the more specific you can be, the better.",
    space_after=10
)

# Table
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"

# Header row
hdr = table.rows[0].cells
hdr[0].text = "Business Profile"
hdr[1].text = "Traction"
hdr[2].text = "Growth"
hdr[3].text = "Soft Landing"
for i, cell in enumerate(hdr):
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1A4E8C')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Early-stage startup, pre-revenue, U.S.-based", "✓", "", ""),
    ("Startup with early revenue, ready to scale", "", "✓", ""),
    ("International company exploring U.S. market entry", "", "", "✓"),
    ("International company ready to establish U.S. operations", "", "", "✓"),
    ("Relocated founder, building from scratch in the U.S.", "✓", "", ""),
    ("Existing company, post-entry, optimizing U.S. operations", "", "✓", ""),
    ("[Other profiles — please add]", "", "", ""),
]

for i, (profile, t, g, sl) in enumerate(rows_data):
    row = table.add_row().cells
    row[0].text = profile
    row[1].text = t
    row[2].text = g
    row[3].text = sl
    bg = 'F2F7FC' if i % 2 == 0 else 'FFFFFF'
    for cell in row:
        for run in cell.paragraphs[0].runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if cell != row[0] else WD_ALIGN_PARAGRAPH.LEFT
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

doc.add_paragraph()
callout(
    "Please review this table. Are the mappings correct? Are there business profiles that don't fit "
    "neatly into one program? Are there profiles we're missing entirely?"
)

# ================================================================
# SECTION 4 — Personas
rule()
heading("4. Learner Personas — Please React", level=2)

body(
    "Three draft personas based on our conversations. "
    "Tell me: who's accurate, who's missing, and which one you see most often in your program.",
    space_after=10
)

personas = [
    {
        "name": "Persona 1 — The Soft Lander",
        "tagline": '"I have a proven business back home. Now I need to make it work here."',
        "desc": (
            "Founder or executive of an established company outside the U.S. — most commonly Latin America "
            "or Europe. The business already works in their home country; they're expanding into the U.S. "
            "and evaluating Central Florida as a base."
        ),
        "bullets": [
            ("Stage:", "Pre-entry or just arrived"),
            ("Biggest challenge:", "Knowing what to do and in what order — legal setup, banking, finding the right people"),
            ("UCF BIP fit:", "Soft Landing"),
        ],
    },
    {
        "name": "Persona 2 — The Latin American Entrepreneur",
        "tagline": '"Central Florida feels like home. But doing business here is different."',
        "desc": (
            "Entrepreneur from a Spanish-speaking country (Venezuela, Colombia, Mexico, Brazil, etc.) "
            "expanding from abroad or recently relocated. Deeply entrepreneurial, but U.S. banking, "
            "legal structures, and professional norms are unfamiliar."
        ),
        "bullets": [
            ("Stage:", "Varies — exploring, newly arrived, or already operating informally"),
            ("Biggest challenge:", "Navigating legal and financial systems in a second language; building a credible professional network"),
            ("UCF BIP fit:", "Soft Landing or Traction"),
        ],
    },
    {
        "name": "Persona 3 — The Relocated Tech Founder",
        "tagline": '"I came here to build. I just don\'t know how the system works yet."',
        "desc": (
            "International entrepreneur who has physically moved to the U.S. — often on an O-1, E-2, "
            "or similar visa. Technically strong and motivated, but overwhelmed by the administrative "
            "side of building a business as a non-citizen."
        ),
        "bullets": [
            ("Stage:", "Pre-revenue to early revenue"),
            ("Biggest challenge:", "IP protection, U.S. contracts, keeping business activities visa-compliant"),
            ("UCF BIP fit:", "Traction or Growth"),
        ],
    },
]

for persona in personas:
    p = doc.add_paragraph()
    r = p.add_run(persona["name"])
    set_font(r, bold=True, size=12)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run(persona["tagline"])
    set_font(r, italic=True, color=(0x2E, 0x75, 0xB6))
    p.paragraph_format.space_after = Pt(6)

    body(persona["desc"], space_after=6)

    for label, value in persona["bullets"]:
        bullet(value, bold_prefix=label + " ")

    doc.add_paragraph()

# ================================================================
# CLOSING
rule()

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("That's it — four things before Monday so we can hit the ground running Thursday:")
set_font(run)

for item in [
    "Outline feedback (Part A and Part B)",
    "Any source materials you can share for Part A",
    "Business profile → program map validation",
    "Persona reactions",
]:
    bullet(item)

doc.add_paragraph()
body("Talk soon.")
doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run("— Gio\n")
set_font(r1, bold=True)
r2 = p.add_run("Design X Factor")
set_font(r2, color=(0x44, 0x44, 0x44))

# Save
output_path = r"G:\z-CUSTOM_DEV\CourseFuture\clients\UCF\BusinessCourse\pre-production-brief-brian.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
