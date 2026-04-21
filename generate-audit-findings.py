#!/usr/bin/env python3
"""
Generate comprehensive content audit findings document.
Covers all feedback sources, sync problems, and issue tracking.
"""
import json, sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
BASE = os.path.dirname(os.path.abspath(__file__))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_DARK  = RGBColor(0x1A, 0x3A, 0x5C)
C_RED   = RGBColor(0xC6, 0x28, 0x28)
C_GREEN = RGBColor(0x2E, 0x7D, 0x32)
C_ORANGE= RGBColor(0xD4, 0x6A, 0x1E)
C_GREY  = RGBColor(0x75, 0x75, 0x75)
C_BLACK = RGBColor(0x1A, 0x1A, 0x1A)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table, color='CCCCCC', size=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for bn in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def pf(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = C_DARK
    return h

def para(doc, text, bold=False, color=C_BLACK, size=10.5, before=0, after=6):
    p = doc.add_paragraph()
    pf(p, before, after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    return p

def bullet(doc, text, color=C_BLACK, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    pf(p, 0, 3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def status_badge(text, color):
    return f'[{text}]'

def add_issue_table(doc, issues):
    """issues = list of (id, lesson, description, source, status, priority)"""
    cols = ['#', 'Lesson', 'Issue', 'Source', 'Status', 'Priority']
    tbl = doc.add_table(rows=1 + len(issues), cols=len(cols))
    tbl.style = 'Table Grid'
    set_table_borders(tbl, '999999', 4)
    widths = [0.3, 0.7, 3.2, 1.2, 0.8, 0.7]

    hrow = tbl.rows[0]
    for i, (h, w) in enumerate(zip(cols, widths)):
        cell = hrow.cells[i]
        set_cell_bg(cell, '1A3A5C')
        cell.width = Inches(w)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    for ri, issue in enumerate(issues):
        trow = tbl.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'F8F8F8'
        for ci, (val, w) in enumerate(zip(issue, widths)):
            cell = trow.cells[ci]
            set_cell_bg(cell, bg)
            cell.width = Inches(w)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(8.5)
            if ci == 4:  # status
                if 'OPEN' in str(val):
                    r.font.color.rgb = C_RED
                    r.font.bold = True
                elif 'DONE' in str(val):
                    r.font.color.rgb = C_GREEN
                elif 'PARTIAL' in str(val):
                    r.font.color.rgb = C_ORANGE
                    r.font.bold = True
            elif ci == 5:  # priority
                if 'P0' in str(val):
                    r.font.color.rgb = C_RED
                    r.font.bold = True
                elif 'P1' in str(val):
                    r.font.color.rgb = C_ORANGE
                    r.font.bold = True
            else:
                r.font.color.rgb = C_BLACK
    doc.add_paragraph()


# ── BUILD DOCUMENT ────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── TITLE ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('UCF Business Immersion Program')
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = C_DARK
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Content Audit Findings — Comprehensive Report')
r2.font.size = Pt(14)
r2.font.color.rgb = C_ORANGE
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('April 7, 2026')
r3.font.size = Pt(11)
r3.font.color.rgb = C_GREY

doc.add_paragraph()

# ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────────
heading(doc, '1. Executive Summary')

para(doc, 'This document reports the findings from a comprehensive audit of ALL feedback sources '
     'for the UCF Business Course content. The audit compared the current JSON content files against '
     'every feedback document, inline editor commit, meeting transcript, and reviewer annotation '
     'collected between February and April 2026.')

para(doc, 'CRITICAL FINDING: Three-Layer Content Sync Problem', bold=True, color=C_RED, size=11)

para(doc, 'The course content exists in THREE separate file layers that have diverged:')
bullet(doc, 'generated/content-phase1-module-N.json (WITH hyphen) — used by the Visualizer/reviewer tool. '
       'Gio and Brian made inline edits here via the CourseFuture Visualizer.')
bullet(doc, 'generated/content-phase1-moduleN.json (NO hyphen) — legacy module files. '
       'Batch update scripts (analogy removal, fun-facts, etc.) operated on these files.')
bullet(doc, 'phase1-lessons/lesson-X-Y.json — source-of-truth lesson files. '
       'Batch scripts also updated these, keeping them in sync with the no-hyphen files.')

para(doc, 'Result: 6 content blocks are OUT OF SYNC between the hyphen files and the other two layers. '
     'The hyphen files contain Gio\'s inline editor refinements AND Brian\'s only direct content edit '
     '(text-2-2-vat in lesson-2-2). Our batch content fixes (analogy removal, "most"\u2192"many", SBDC URLs, '
     'fun-facts) were applied to the no-hyphen and phase1-lessons files but NOT to the hyphen files.',
     color=C_RED)

para(doc, 'This means: (1) Brian\'s direct edit was never propagated, and (2) our quality fixes '
     'were never applied to the files the Visualizer displays to Brian.', bold=True)

doc.add_paragraph()

# ── 2. FEEDBACK SOURCES INVENTORY ────────────────────────────────────────
heading(doc, '2. Feedback Sources Inventory')

sources_data = [
    ('Inline Editor Commits', 'Apr 1\u20133', '10 commits', 'Gio edited lesson-0-1 (7 blocks), lesson-1-1 (1 block), '
     'lesson-2-2 (2 blocks). Brian edited text-2-2-vat in lesson-2-2.'),
    ('April 2 Meeting Transcript', 'Apr 2', '62 KB transcript', 'Video strategy pivot, navigation rework, '
     'I-4/on-ramp too specific, avatar voice feedback, content/UX decisions.'),
    ('April 2 Meeting Analysis', 'Apr 2', '68 KB .docx', 'Structured analysis of meeting: video format (bookend), '
     'action items, timeline, risk flags.'),
    ('April 6 Themes_feedback.docx', 'Apr 6', '8.8 MB .docx', 'Brian\'s location name choices for navigation: '
     'The Downtown, Financial District, Space Center, Theme Park, The Port, The Beach, The Springs, UCF. '
     'Alternatives: Hotels, Orange Grooves, Shops, Stadiums.'),
    ('April 6 UCF-Video-Scripts-Draft.docx', 'Apr 6', '46 KB .docx', 'Full video scripts for all 9 modules '
     '(INTRO + BRIDGE + RECAP format). Spanish translation issues flagged. Awaiting Brian review.'),
    ('Mar 19 feedback.txt', 'Mar 19', '3.2 KB email', 'Tone: absolute language ("most"\u2192"many"), '
     'reflection checkpoints too many questions, "their business"\u2192"your business".'),
    ('Mar 19 Module-1 .docx', 'Mar 19', '35 KB .docx', 'Module 1 review with inline comments: '
     'IRS phone line, domain name before SunBiz, knowledge check reconsidering, redundancy.'),
    ('Mar 18 Module-0 .docx', 'Mar 18', '27 KB .docx', 'Module 0 review: combine themes+pitfalls into table, '
     'proposed knowledge check format, "road trip" → "journey" language.'),
    ('Mar 18 Meeting Transcript', 'Mar 18', '42 KB transcript', 'Discussion of content direction, '
     'review process, editorial nuances.'),
]

tbl = doc.add_table(rows=1 + len(sources_data), cols=4)
tbl.style = 'Table Grid'
set_table_borders(tbl, '999999', 4)
hrow = tbl.rows[0]
for i, h in enumerate(['Source', 'Date', 'Size', 'Key Content']):
    cell = hrow.cells[i]
    set_cell_bg(cell, '1A3A5C')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

for ri, row in enumerate(sources_data):
    trow = tbl.rows[ri + 1]
    bg = 'FFFFFF' if ri % 2 == 0 else 'F8F8F8'
    for ci, val in enumerate(row):
        cell = trow.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(8.5)
        r.font.color.rgb = C_BLACK

doc.add_paragraph()

# ── 3. SYNC PROBLEM DETAIL ───────────────────────────────────────────────
heading(doc, '3. Critical: Content Sync Problem — 6 Blocks Out of Sync')

para(doc, 'The sync comparison ran across all 9 modules (26 lessons, ~300 content blocks). '
     '6 blocks have content that differs between the hyphen file (Visualizer) and the '
     'no-hyphen/phase1-lessons files (batch scripts, HTML generator).', before=4)

sync_issues = [
    ('lesson-0-1', 'text-0-1-0', 'Opening hook',
     'Hyphen: "not paperwork\u2014but what you don\'t know yet? Most founders..." '
     'No-hyphen/phase1: "not what you think it is? Many founders..." '
     'Our "Most\u2192Many" fix was only applied to no-hyphen/phase1.'),
    ('lesson-0-1', 'text-0-1-1', '"Expanding into the U.S."',
     'Hyphen: Gio\'s rewrite \u2014 "not just a series of tasks to complete" + "jump in\u2014form a company" '
     'No-hyphen/phase1: Original \u2014 "not just ticking boxes" + "dive in, form an LLC"'),
    ('lesson-0-1', 'text-0-1-learn-first', 'Program intro paragraph',
     'Hyphen: Gio\'s rewrite \u2014 "designed as a low-commitment first step" + BIP continuation mention '
     'No-hyphen/phase1: Original \u2014 "the first step in getting ready for the U.S. market"'),
    ('lesson-0-1', 'text-0-1-3', 'How to Use This Program',
     'Hyphen: Clean standalone-modules format, no "self-paced program" prefix '
     'No-hyphen/phase1: Has duplicate "self-paced program..." paragraph prepended by batch fix'),
    ('lesson-0-1', 'kt-0-1-1', 'Key takeaway',
     'Hyphen: "You haven\'t entered the U.S. market yet\u2014you\'ve mapped it." (Gio inline edit) '
     'No-hyphen/phase1: "You have not entered the U.S. market yet \u2014 you have mapped it." (our batch fix)'),
    ('lesson-2-2', 'text-2-2-vat', 'VAT vs Sales Tax intro',
     'Hyphen: Brian\'s edit \u2014 "you may be familiar with" + removed Financial District analogy '
     'No-hyphen/phase1: Original \u2014 "you\'re familiar with" + still has Financial District analogy. '
     'THIS IS BRIAN\'S ONLY DIRECT CONTENT EDIT AND IT WAS NEVER PROPAGATED.'),
]

tbl2 = doc.add_table(rows=1 + len(sync_issues), cols=4)
tbl2.style = 'Table Grid'
set_table_borders(tbl2, '999999', 4)
hrow2 = tbl2.rows[0]
for i, h in enumerate(['Lesson', 'Block ID', 'Description', 'Divergence Detail']):
    cell = hrow2.cells[i]
    set_cell_bg(cell, 'C62828')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

for ri, (lesson, block, desc, detail) in enumerate(sync_issues):
    trow = tbl2.rows[ri + 1]
    bg = 'FFF3E0' if ri % 2 == 0 else 'FFECB3'
    for ci, val in enumerate([lesson, block, desc, detail]):
        cell = trow.cells[ci]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(8)
        r.font.color.rgb = C_BLACK
        if ci == 1:
            r.font.bold = True

doc.add_paragraph()
para(doc, 'RESOLUTION: Before any content regeneration, all 6 blocks must be reconciled. '
     'The hyphen file versions (which include Gio\'s inline refinements and Brian\'s edit) should be '
     'treated as the authoritative source and propagated to the no-hyphen and phase1-lessons files. '
     'Then our batch fixes (most\u2192many, analogy removal, fun-facts, SBDC URLs) must be re-applied '
     'to the merged content.', bold=True, color=C_RED)

# ── 4. COMPLETE ISSUE TRACKING TABLE ─────────────────────────────────────
heading(doc, '4. Complete Issue Tracking — All Feedback Items')

para(doc, 'Every discrete feedback item from all sources, with current implementation status. '
     'P0 = must fix before Brian review. P1 = fix before delivery. P2 = nice to have.')

all_issues = [
    # Sync issues
    ('S1', '0-1', 'text-0-1-0 diverged: hyphen has different opening hook than no-hyphen/phase1', 'Sync gap', 'OPEN', 'P0'),
    ('S2', '0-1', 'text-0-1-1 diverged: hyphen has Gio\'s rewrite of "expanding" paragraph', 'Sync gap', 'OPEN', 'P0'),
    ('S3', '0-1', 'text-0-1-learn-first diverged: hyphen has simplified program intro', 'Sync gap', 'OPEN', 'P0'),
    ('S4', '0-1', 'text-0-1-3 diverged: hyphen has clean "how to use" text; no-hyphen has buggy duplicate', 'Sync gap', 'OPEN', 'P0'),
    ('S5', '0-1', 'kt-0-1-1 diverged: minor wording differences between hyphen and no-hyphen', 'Sync gap', 'OPEN', 'P0'),
    ('S6', '2-2', 'text-2-2-vat: Brian\'s edit ("you may be familiar" + removed analogy) never propagated', 'Sync gap', 'OPEN', 'P0'),

    # April 2 meeting
    ('A1', 'All', 'Navigation: lead with LEARNING TOPIC not location name (Brian: international users don\'t know FL places)', 'Apr 2 meeting', 'OPEN', 'P1'),
    ('A2', '0-1', 'I-4 / on-ramp references too specific for international audience — make generic', 'Apr 2 meeting', 'PARTIAL', 'P1'),
    ('A3', 'All', 'Avatar face (floating head image) in quiz/scenario feedback sections', 'Apr 2 meeting', 'OPEN', 'P2'),
    ('A4', 'Video', 'Female voice needs more Latina accent (male voice approved)', 'Apr 2 meeting', 'OPEN', 'P1'),
    ('A5', '0-1', 'Intro video should be generic, not lesson-0-1; develop LAST', 'Apr 2 meeting', 'OPEN', 'P2'),
    ('A6', 'Video', 'Video strategy pivot: bookend format (INTRO ~60s + BRIDGE ~15-20s + RECAP ~20-30s)', 'Apr 2 meeting', 'DONE', 'P0'),
    ('A7', 'Video', 'Video scripts written in draft form, awaiting Brian review', 'Apr 2 meeting', 'OPEN', 'P1'),
    ('A8', 'All', 'Hyperlinks to resources confirmed: link directly when a resource is mentioned', 'Apr 2 meeting', 'DONE', 'P0'),
    ('A9', 'All', 'Reflection textareas removed (Brian agreed: "meaningless")', 'Apr 2 meeting', 'DONE', 'P0'),
    ('A10', '8-x', 'Module 8 should signal conclusion ("End of the Road" or similar)', 'Apr 2 meeting', 'OPEN', 'P2'),

    # April 6 feedback
    ('T1', 'Nav', 'Brian\'s location names: The Downtown, Financial District, Space Center, Theme Park, The Port, The Beach, The Springs, UCF', 'Apr 6 Themes', 'OPEN', 'P1'),
    ('T2', 'Nav', 'Alternative location names suggested: Hotels, Orange Grooves, Shops, Stadiums — decision pending', 'Apr 6 Themes', 'OPEN', 'P2'),
    ('T3', 'Video', 'Spanish translation issues in video scripts: "hit the road", "deal went cold", "unwritten rules" idioms', 'Apr 6 Scripts', 'OPEN', 'P1'),

    # Mar 19 feedback (previously addressed)
    ('M1', '1-1', '"their visa or tax situation" \u2192 "your visa..." (address user directly)', 'Mar 19', 'DONE', 'P0'),
    ('M2', 'All', '"most" \u2192 "many" for absolute language (M0, M1)', 'Mar 19', 'DONE', 'P0'),
    ('M3', '1-3', '"Orlando-Miami corridor" \u2192 "Central Florida"', 'Mar 19', 'DONE', 'P0'),
    ('M4', '0-1', '"Orlando area" \u2192 "Central Florida"', 'Mar 19', 'DONE', 'P0'),
    ('M5', '1-1', 'C-Corp conversion story: soften specific dollar/time amounts and attributed quote', 'Mar 19', 'DONE', 'P0'),

    # Mar 19 feedback (NOT yet addressed)
    ('M6', '1-2', 'IRS phone line: "field test that IRS is actually answering this line" + wait times', 'Mar 19 comment', 'OPEN', 'P1'),
    ('M7', '1-1', 'Knowledge check: reconsidering Q \u2014 "options are variations of only 2 options (DIY or pay)"', 'Mar 19 comment', 'OPEN', 'P1'),
    ('M8', '1-1', 'LLC to C-Corp conversion story told TWICE \u2014 check for redundancy', 'Mar 19 overall', 'OPEN', 'P1'),
    ('M9', 'All', '"Central Florida" not "Orlando" or "Orlando area" throughout', 'Mar 19 overall', 'DONE', 'P0'),
    ('M10', '1-2', 'Add domain name check recommendation before finalizing SunBiz', 'Mar 19 comment', 'DONE', 'P0'),
    ('M11', '1-x', 'LLC table: "any visa holder can own" \u2192 "Ownership is generally allowed, but your visa determines whether you can actively work in the business"', 'Mar 19 / Apr 3', 'OPEN', 'P1'),

    # Mar 18 feedback (Module 0)
    ('O1', '0-1', 'Combine themes + common pitfalls into single table/1-pager (Brian\'s idea)', 'Mar 18 M0', 'OPEN', 'P2'),
    ('O2', '0-1', 'Disclaimer: exact wording per Brian \u2014 "does not constitute legal, tax, immigration, or financial advice"', 'Mar 18 M0 / Apr 3', 'OPEN', 'P1'),
    ('O3', '0-1', '"How to Use" section: rewrite as narrative, not checklist', 'Mar 18 M0 / Apr 3', 'PARTIAL', 'P1'),
    ('O4', '0-1', 'Common Mistakes callout: trim to 1\u20132 examples max, tell users to look for it in each module', 'Apr 3 comment', 'OPEN', 'P1'),
    ('O5', '0-1', 'Be explicit: 8 content modules (1\u20138) + 1 intro (Module 0) = 9 total', 'Apr 3 comment', 'OPEN', 'P1'),

    # Content quality (from our gate runs)
    ('Q1', 'All', 'Fun-fact callout added to every lesson (Florida location \u2192 topic)', 'Content QA', 'DONE', 'P0'),
    ('Q2', 'All', 'Road trip / journey analogy removed from all body text (retained in fun-facts only)', 'Content QA', 'DONE', 'P0'),
    ('Q3', 'All', '"3\u20136 months" \u2192 "2\u20136 months" timeline references', 'Content QA', 'DONE', 'P0'),
    ('Q4', 'All', 'SBDC URL standardized to sbdcorlando.com', 'Content QA', 'DONE', 'P0'),
    ('Q5', '7-3', 'Lesson-7-3 depth: added scenario "Reading the Numbers" for economic data application', 'Gate 1.9', 'DONE', 'P1'),
    ('Q6', 'All', 'Module 3 location: "Space Center" / "El Centro Espacial" (not Launch Pad or Cape Canaveral)', 'Content QA', 'DONE', 'P0'),
    ('Q7', 'All', 'Spanish translations need updating for all English content changes made', 'Bilingual', 'OPEN', 'P1'),
]

add_issue_table(doc, all_issues)

# ── 5. SUMMARY COUNTS ────────────────────────────────────────────────────
heading(doc, '5. Issue Summary by Status')

counts = {'OPEN': 0, 'PARTIAL': 0, 'DONE': 0}
p0_open = 0
p1_open = 0
for issue in all_issues:
    status = issue[4]
    counts[status] = counts.get(status, 0) + 1
    if status in ('OPEN', 'PARTIAL') and issue[5] == 'P0':
        p0_open += 1
    if status in ('OPEN', 'PARTIAL') and issue[5] == 'P1':
        p1_open += 1

para(doc, f'Total issues tracked: {len(all_issues)}', bold=True)
bullet(doc, f'DONE: {counts["DONE"]} items (implemented and verified)', C_GREEN)
bullet(doc, f'OPEN: {counts["OPEN"]} items (not yet addressed)', C_RED)
bullet(doc, f'PARTIAL: {counts["PARTIAL"]} items (partially addressed)', C_ORANGE)
bullet(doc, f'P0 (critical) still open: {p0_open}', C_RED)
bullet(doc, f'P1 (before delivery) still open: {p1_open}', C_ORANGE)

doc.add_paragraph()

# ── 6. RECOMMENDED ACTIONS ────────────────────────────────────────────────
heading(doc, '6. Recommended Actions for Content Regeneration')

para(doc, 'STEP 1 \u2014 Resolve Sync Problem (P0, blocks all other work)', bold=True, color=C_RED, size=11)
bullet(doc, 'Merge hyphen file content into no-hyphen and phase1-lessons files')
bullet(doc, 'For lesson-0-1: Use the hyphen file versions (Gio\'s inline refinements) as the base, '
       'then re-apply our batch fixes on top (most\u2192many, fun-fact wording, SBDC URL)')
bullet(doc, 'For lesson-2-2: Propagate Brian\'s text-2-2-vat edit (hedged language + removed analogy) '
       'to no-hyphen and phase1-lessons')
bullet(doc, 'After merge: run the full sync comparison again to verify 0 mismatches')

para(doc, 'STEP 2 \u2014 Apply Remaining Mar 19 Feedback (P1)', bold=True, color=C_ORANGE, size=11)
bullet(doc, 'M6: Verify IRS phone number is current; add estimated wait time')
bullet(doc, 'M7: Redesign lesson-1-1 knowledge check (too binary)')
bullet(doc, 'M8: Check if LLC\u2192C-Corp conversion story appears more than once')
bullet(doc, 'M11: Fix LLC table cell in lesson-1-1 comparison card')

para(doc, 'STEP 3 \u2014 Apply Remaining Module 0 Feedback (P1)', bold=True, color=C_ORANGE, size=11)
bullet(doc, 'O2: Verify disclaimer uses Brian\'s exact wording')
bullet(doc, 'O3: Finalize "How to Use" section as narrative')
bullet(doc, 'O4: Trim Common Mistakes callout to 1\u20132 examples')
bullet(doc, 'O5: Make "8 content modules + 1 intro = 9 total" explicit')

para(doc, 'STEP 4 \u2014 Navigation Rework (P1)', bold=True, color=C_ORANGE, size=11)
bullet(doc, 'Reconcile Apr 2 meeting ("lead with learning topic") vs Apr 6 Themes_feedback (location names)')
bullet(doc, 'Implement Brian\'s location choices: The Downtown, Financial District, Space Center, etc.')
bullet(doc, 'Update index.html navigation with finalized names + descriptions')

para(doc, 'STEP 5 \u2014 Spanish Translation Update (P1)', bold=True, color=C_ORANGE, size=11)
bullet(doc, 'Update all lesson-*-es.json overlay files to match English content changes')
bullet(doc, 'Address video script idiom translation issues before voice generation')

para(doc, 'STEP 6 \u2014 Video Pipeline (P1, after Steps 1\u20135)', bold=True, color=C_ORANGE, size=11)
bullet(doc, 'Get Brian\'s approval on video scripts (April 6 draft document)')
bullet(doc, 'Source/configure Latina accent female voice')
bullet(doc, 'Generate voiceover audio for Brian review BEFORE any video rendering')

para(doc, 'STEP 7 \u2014 Nice-to-Have Enhancements (P2)', bold=True, color=C_GREY, size=11)
bullet(doc, 'A3: Avatar floating head in quiz feedback')
bullet(doc, 'O1: Combined themes/pitfalls 1-pager for Module 0')
bullet(doc, 'A5: Generic intro video (separate from lesson-0-1)')
bullet(doc, 'T2: Finalize alternative location names with Brian')

doc.add_paragraph()

# ── 7. KEY BRIAN QUOTES ──────────────────────────────────────────────────
heading(doc, '7. Key Brian Quotes (for editorial alignment)')

quotes = [
    ('On absolute language', '"Be particular about using words like \'most\' when \'many\' might be more accurate or less controversial. Keep a risk and liability lens on statements."'),
    ('On case study specifics', '"Let\'s consider what level of detail we use. It may be safer to use less specific dollar and time ranges."'),
    ('On user address', '"Please change the language throughout to address the user directly such as \'your business\' (not \'their business\')."'),
    ('On geography', '"Try to stick with Central Florida rather than Orlando or the Orlando area." / "Some content references Orlando-Miami corridor. Avoid that."'),
    ('On video format', '"I\'m biased because I\'m not a video learner. I\'m very text driven." / Videos should be "like a TV host who introduces something, they play the thing, and then they do color commentary."'),
    ('On navigation', '"Only we three understand what Florida is. Somebody else coming is like, okay, you know."'),
    ('On review process', '"I will like to make several edits in this category before the content is finalized."'),
    ('On redundancy', '"Check for redundancy. The investor who had to convert from LLC to C-Corp story is told twice."'),
    ('On reflection prompts', '"The reflections have a lot of questions. That makes it comprehensive, but may be overwhelming."'),
]

for title, quote in quotes:
    p = doc.add_paragraph()
    pf(p, 4, 2)
    r1 = p.add_run(f'{title}: ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = C_DARK
    r2 = p.add_run(quote)
    r2.font.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = C_GREY

# ── SAVE ──────────────────────────────────────────────────────────────────
out = os.path.join(BASE, 'UCF-Content-Audit-Findings.docx')
doc.save(out)
print(f'\n\u2713 Saved: UCF-Content-Audit-Findings.docx')
