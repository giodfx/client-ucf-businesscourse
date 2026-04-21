import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after  = Pt(4)
    return h

def add_para(doc, text, bold=False, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def style_cell(cell, text, bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_table_header_row(table, cols, bg='1A3A5C'):
    row = table.rows[0]
    for i, col in enumerate(cols):
        style_cell(row.cells[i], col, bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color='FFFFFF')
        set_cell_bg(row.cells[i], bg)

# ── TITLE PAGE ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('U.S. Market Readiness Program')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Navigation & Content Recommendations')
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Prepared for UCF Business Incubation Program  ·  April 2026')
r3.font.size = Pt(10)
r3.italic = True
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ── PURPOSE ───────────────────────────────────────────────────────────────────
add_heading(doc, '1. Purpose of This Document', 1)
add_para(doc,
    'This document consolidates recommendations for the course navigation structure, '
    'module naming conventions, and video script revisions for the U.S. Market Readiness Program. '
    'It covers both English and Spanish text and is intended for review and approval '
    'by the UCF BIP team before implementation.')

# ── DESIGN FRAMEWORK ─────────────────────────────────────────────────────────
add_heading(doc, '2. Design Framework — The Road Trip Analogy', 1)
add_para(doc,
    'The course uses a road trip through Central Florida as its visual and narrative theme. '
    'This framework is effective when applied at the right layer of the experience.')

add_heading(doc, 'Where the analogy works', 2)
add_bullet(doc, 'The dashboard interface — the car, highway, and Florida landscape backgrounds')
add_bullet(doc, 'Avatar video transitions that move learners between module locations')
add_bullet(doc, 'Interface language: "Destinations," "Select Your Destination"')
add_bullet(doc, 'Module location names as atmospheric, visual anchors')

add_heading(doc, 'Where it creates confusion', 2)
add_bullet(doc, 'When Florida-specific location names are the primary navigation label with no topic context')
add_bullet(doc, 'When idioms or expressions require local knowledge to understand (e.g., "Getting Your Foot in the Door")')
add_bullet(doc, 'When location names do not tell the learner what they will find inside the module')

add_heading(doc, 'The guiding principle', 2)
add_para(doc,
    'Keep the road trip visual experience fully intact. The landscape images, dashboard, and avatar '
    'videos carry the atmosphere. Navigation text should be functional and clear — '
    'it tells learners exactly where to stop and why, in plain language that translates '
    'cleanly into Spanish without losing meaning.')

# ── MODULE NAMES ─────────────────────────────────────────────────────────────
add_heading(doc, '3. Recommended Module Names', 1)
add_para(doc,
    'Each module uses a two-part label: the Florida location name (atmosphere) '
    'and a sentence description (what the learner will be able to do after completing it). '
    'Both are provided in English and Spanish.')

doc.add_paragraph()

# Table
tbl = doc.add_table(rows=10, cols=4)
tbl.style = 'Table Grid'
tbl.autofit = False

# Column widths
col_widths = [1.1, 1.5, 2.5, 2.5]
for i, w in enumerate(col_widths):
    for row in tbl.rows:
        row.cells[i].width = Inches(w)

# Header
add_table_header_row(tbl, ['Module', 'Location', 'Description (EN)', 'Descripción (ES)'])

modules = [
    ('Start Here',          'Highway\nLa Autopista',
     'Get clear on what it actually takes to enter the U.S. market',
     'Entienda lo que realmente implica entrar al mercado de EE.UU.'),
    ('Module 1',            'Downtown\nEl Centro',
     'Learn how to structure, register, and open your U.S. business',
     'Aprenda a estructurar, registrar y abrir su negocio en EE.UU.'),
    ('Module 2',            'Financial District\nEl Distrito Financiero',
     'Learn how the U.S. tax and banking system works for your business',
     'Entienda cómo funciona el sistema fiscal y bancario de EE.UU.'),
    ('Module 3',            'Launch Pad\nLa Plataforma de Lanzamiento',
     'Shield your business with the right IP, insurance, and contracts',
     'Proteja su negocio con la propiedad intelectual, seguros y contratos correctos'),
    ('Module 4',            'Theme Park District\nLos Parques Temáticos',
     'Build your U.S. team without costly legal and compliance mistakes',
     'Forme su equipo en EE.UU. sin cometer errores legales costosos'),
    ('Module 5',            'Harbor\nEl Puerto',
     'Learn how to find, reach, and close your first U.S. customers',
     'Aprenda a encontrar, llegar y vender a sus primeros clientes en EE.UU.'),
    ('Module 6',            'Beach\nLa Playa',
     'Discover how Americans communicate, negotiate, and make decisions',
     'Descubra cómo los americanos se comunican, negocian y deciden'),
    ('Module 7',            'Natural Springs\nLos Manantiales',
     'Connect with the free resources most international founders overlook',
     'Conéctese con los recursos gratuitos que la mayoría de los fundadores ignoran'),
    ('Module 8',            'UCF Campus\nEl Campus UCF',
     'Assess your readiness and plan your next steps with UCF BIP',
     'Evalúe su preparación y planifique sus próximos pasos con UCF BIP'),
]

for i, (mod, loc, desc_en, desc_es) in enumerate(modules):
    row = tbl.rows[i + 1]
    style_cell(row.cells[0], mod,     bold=True,  size=9)
    style_cell(row.cells[1], loc,     bold=False, size=9)
    style_cell(row.cells[2], desc_en, bold=False, size=9)
    style_cell(row.cells[3], desc_es, bold=False, size=9, italic=True)
    bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
    for c in row.cells:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ── INTERFACE TEXT ────────────────────────────────────────────────────────────
add_heading(doc, '4. Navigation Interface Text', 1)
add_para(doc, 'The following text elements appear in the course interface. Recommended final values:')

# Interface table
itbl = doc.add_table(rows=7, cols=3)
itbl.style = 'Table Grid'
itbl.autofit = False
itbl.columns[0].width = Inches(2.0)
itbl.columns[1].width = Inches(2.5)
itbl.columns[2].width = Inches(2.5)

add_table_header_row(itbl, ['Element', 'English', 'Spanish'])

irows = [
    ('Course badge (top bar)',    'U.S. Market Readiness Program',        'Programa de Preparación para el Mercado de EE.UU.'),
    ('Dashboard button',          'DESTINATIONS',                          'DESTINOS'),
    ('Navigation panel title',    'Select Your Destination',               'Seleccione Su Destino'),
    ('Onboarding Step 1 title',   'Welcome',                               'Bienvenido'),
    ('Onboarding Step 2 title',   'Before You Begin',                      'Antes de Empezar'),
    ('Onboarding Step 3 title',   'Important Disclaimer',                  'Aviso Importante'),
]

for i, (elem, en, es) in enumerate(irows):
    row = itbl.rows[i + 1]
    style_cell(row.cells[0], elem, bold=True,  size=9)
    style_cell(row.cells[1], en,   bold=False, size=9)
    style_cell(row.cells[2], es,   bold=False, size=9, italic=True)
    bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
    for c in row.cells:
        set_cell_bg(c, bg)

doc.add_paragraph()

# ── VIDEO SCRIPT ISSUES ───────────────────────────────────────────────────────
add_heading(doc, '5. Video Script — Items to Resolve Before Production', 1)
add_para(doc,
    'The following issues were identified in the Preliminary Draft (April 2, 2026). '
    'These must be resolved before voiceover audio is generated.')

add_heading(doc, '5.1  Lesson 7-5 Is Missing from the Scripts  [CRITICAL]', 2)
add_para(doc,
    'The scripts cover Module 7 as having 4 lessons (7-1 through 7-4). '
    'The course contains 5 lessons in Module 7. Lesson 7-5 — "UCF BIP Programs — '
    'Your Next Step" — is a full 6-minute lesson with a knowledge check, application '
    'readiness checklist, and a UCF BIP roadmap exercise. It is not covered by any video.')
add_bullet(doc, 'A BRIDGE video from Lesson 7-4 to Lesson 7-5 is missing entirely')
add_bullet(doc, 'The RECAP video is placed at Lesson 7-4 but should be at Lesson 7-5 (the final lesson of Module 7)')
add_bullet(doc, '',
           bold_prefix='Note: ')
# fix that bullet
p = doc.paragraphs[-1]
p.clear()
p.style = doc.styles['List Bullet']
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Note: ')
r.bold = True
p.add_run(
    'Lesson 7-5 and Lesson 8-3 ("Next Steps — Phases 2 & 3") cover similar content — '
    'both describe UCF BIP Phases 2 and 3. Confirm whether this overlap is intentional reinforcement '
    'or whether one should be consolidated.')

add_heading(doc, '5.2  "Competitive Analysis Framework" Does Not Exist  [MUST FIX]', 2)
add_para(doc,
    'The Module 5 INTRO script says: "The competitive analysis framework in this module '
    'is one you\'ll use repeatedly." No such tool exists in Lesson 5-1. '
    'The lesson contains a pricing scenario and a knowledge check, but no interactive framework tool.')
add_bullet(doc, 'Either build and add the tool to Lesson 5-1, or rewrite the script line to describe what is actually there')

add_heading(doc, '5.3  "Entity Comparison Tool" Terminology', 2)
add_para(doc,
    'The Module 1 INTRO refers to "the entity comparison tool." '
    'The actual component in Lesson 1-1 is a branching decision tree — '
    'it asks about your situation and recommends LLC, C-Corp, or S-Corp. '
    'It is not a side-by-side comparison chart.')
add_bullet(doc, 'Suggested fix: change "entity comparison tool" to "entity decision guide" in the script')

add_heading(doc, '5.4  Spanish Idioms Flagged in General Notes — Not Yet Resolved', 2)
add_para(doc,
    'The document\'s own General Notes section flagged the following phrases as '
    'potentially problematic for Spanish translation. All nine still appear in the scripts unchanged:')

idioms = [
    ('"Hit the road"',                       'Module 0 INTRO closing line',   'Replace with: "Let\'s get started" / "Comencemos"'),
    ('"Deal went cold"',                      'Module 6 INTRO',                'Replace with: "a promising deal fell through" / "un trato prometedor no se concretó"'),
    ('"Runs on great teams behind the scenes"','Module 4 INTRO',               'Replace with: "depends on well-managed teams" / "depende de equipos bien gestionados"'),
    ('"Unwritten rules"',                     'Module 6 INTRO',                'Acceptable if explained in context; otherwise: "norms that no one explains openly"'),
    ('"What to watch for"',                   'Module 2 INTRO (repeated)',     'Replace with: "What you need to know" / "Lo que debe saber"'),
    ('"Sales tax is nothing like VAT"',       'Module 2 INTRO',                'Replace with: "U.S. sales tax works very differently from VAT"'),
    ('"Rocket base"',                         'Module 3 location reference',   'Use "launch pad" or "space center" — both translate clearly'),
    ('"Theme park district"',                 'Module 4 location reference',   'Use "Theme Park District" as location label; avoid it as a descriptor in spoken script'),
    ('"Doesn\'t do it alone"',               'Module 7 INTRO',                'Acceptable — clear enough; minor simplification only if needed'),
]

itbl2 = doc.add_table(rows=len(idioms)+1, cols=3)
itbl2.style = 'Table Grid'
itbl2.autofit = False
itbl2.columns[0].width = Inches(1.6)
itbl2.columns[1].width = Inches(1.5)
itbl2.columns[2].width = Inches(3.0)
add_table_header_row(itbl2, ['Phrase', 'Location in Scripts', 'Recommended Fix'])
for i, (phrase, loc, fix) in enumerate(idioms):
    row = itbl2.rows[i+1]
    style_cell(row.cells[0], phrase, bold=True,  size=8.5, italic=True)
    style_cell(row.cells[1], loc,    bold=False, size=8.5)
    style_cell(row.cells[2], fix,    bold=False, size=8.5)
    bg = 'F0F4F8' if i % 2 == 0 else 'FFFFFF'
    for c in row.cells:
        set_cell_bg(c, bg)

doc.add_paragraph()

add_heading(doc, '5.5  Module 0 INTRO Duration', 2)
add_para(doc,
    'The format note states INTRO videos should be approximately 60 seconds. '
    'The Module 0 INTRO is labeled "~90 seconds." This is likely appropriate '
    'for the course introduction, but should be noted as an intentional exception '
    'so it is not cut during production.')

# ── SPANISH PRINCIPLES ────────────────────────────────────────────────────────
add_heading(doc, '6. Spanish Translation Principles', 1)
add_para(doc,
    'The following principles apply to all Spanish content in this course '
    '— navigation labels, video scripts, lesson text, and interface elements.')

add_bullet(doc, '',
           bold_prefix='Plain language first. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Bullet']
p.paragraph_format.space_after = Pt(4)
p.add_run('Plain language first.  ').bold = True
p.add_run('If a phrase requires cultural knowledge of the United States or Central Florida to understand, it will not translate well. Write for a founder in Bogotá, São Paulo, or Mexico City who has never visited Florida.')

add_bullet(doc, '',
           bold_prefix='Avoid action idioms. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Bullet']
p.paragraph_format.space_after = Pt(4)
p.add_run('Avoid action idioms.  ').bold = True
p.add_run('Phrases like "hit the road," "get your foot in the door," or "watch the money" do not translate literally. Use direct verbs: start, register, protect, hire, reach.')

add_bullet(doc, '',
           bold_prefix='Verb-led sentences translate best. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Bullet']
p.paragraph_format.space_after = Pt(4)
p.add_run('Verb-led sentences translate best.  ').bold = True
p.add_run('"Learn how the U.S. tax system works" → "Aprenda cómo funciona el sistema fiscal de EE.UU." — direct, clear, identical meaning in both languages.')

add_bullet(doc, '',
           bold_prefix='Use neutral Spanish. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Bullet']
p.paragraph_format.space_after = Pt(4)
p.add_run('Use neutral Spanish.  ').bold = True
p.add_run('Avoid regional expressions from any single country. The audience is pan-Latin American. Terms like "Entienda," "Aprenda," "Proteja," "Conéctese" are universally understood formal commands appropriate for a business education context.')

add_bullet(doc, '',
           bold_prefix='Location names: keep them simple. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Bullet']
p.paragraph_format.space_after = Pt(4)
p.add_run('Location names: keep them simple.  ').bold = True
p.add_run('"El Centro," "El Distrito Financiero," "El Puerto," "La Playa" — these are universally understood across all Spanish-speaking countries. Avoid over-localizing (e.g., "Cabo Cañaveral" is fine as a geographic name but should not appear as the primary navigation label).')

# ── NEXT STEPS ────────────────────────────────────────────────────────────────
add_heading(doc, '7. Recommended Next Steps', 1)

add_numbered(doc, '', bold_prefix='Review and approve module names. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Number']
p.paragraph_format.space_after = Pt(5)
p.add_run('Review and approve module names.  ').bold = True
p.add_run('Confirm the location names and sentence descriptions in Section 3. Mark any changes directly in this document.')

add_numbered(doc, '', bold_prefix='Resolve video script issues. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Number']
p.paragraph_format.space_after = Pt(5)
p.add_run('Resolve video script issues.  ').bold = True
p.add_run('Address the five items in Section 5, particularly the missing Lesson 7-5 coverage and the competitive analysis framework reference. An updated script draft should be produced before voiceover recording begins.')

add_numbered(doc, '', bold_prefix='Confirm Lesson 7-5 vs. Lesson 8-3. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Number']
p.paragraph_format.space_after = Pt(5)
p.add_run('Confirm Lesson 7-5 vs. Lesson 8-3.  ').bold = True
p.add_run('Both lessons cover UCF BIP Phases 2 and 3. Confirm whether both remain as separate lessons or whether the content should be consolidated.')

add_numbered(doc, '', bold_prefix='Implement navigation changes. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Number']
p.paragraph_format.space_after = Pt(5)
p.add_run('Implement navigation changes.  ').bold = True
p.add_run('Once module names are approved, update the course dashboard, lesson headers, and any promotional or onboarding materials to reflect the final naming.')

add_numbered(doc, '', bold_prefix='Provide case studies for Lesson 7-5. ')
p = doc.paragraphs[-1]; p.clear(); p.style = doc.styles['List Number']
p.paragraph_format.space_after = Pt(5)
p.add_run('Provide case studies for Lesson 7-5.  ').bold = True
p.add_run('Three company case study placeholders are currently in the lesson. Brian to supply company details by the agreed date.')

# ── Save ──────────────────────────────────────────────────────────────────────
out = 'G:/z-CUSTOM_DEV/CourseFuture/clients/UCF/BusinessCourse/UCF-Navigation-Recommendations.docx'
doc.save(out)
print('Saved to:', out)
